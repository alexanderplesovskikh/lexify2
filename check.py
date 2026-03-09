#imports
import os
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
import re
from habanero import cn
from citeproc.source.bibtex import BibTeX
from citeproc import CitationStylesStyle, CitationStylesBibliography
from citeproc import formatter
from citeproc import Citation, CitationItem
from io import StringIO, BytesIO
import sys
import subprocess
import fitz
from xml.etree import ElementTree as ET
from docx.oxml import parse_xml

#from __future__ import annotations

from dataclasses import dataclass
from typing import Any, Optional
import os
import re
import requests

CROSSREF_WORKS_URL = "https://api.crossref.org/works"

#params
translate_alignmnet_options = {
    "left": "по левому краю",
    "right": "по правому краю",
    "center": "по центру",
    "both": "по ширине",
    "top": "сверху",
    "bottom": "снизу",
}

translate_table_options = {
    "left": "по левому краю",
    "right": "по правому краю",
    "center": "по центру",
    "both": "по ширине",
    "top": "сверху",
    "bottom": "снизу",
}

ref_keyword_list = list(set(['списокиспользованныхисточников', 
                    'списокиспользуемыхисточников', 
                    'списокисточников', 
                    'списоклитературы', 
                    'публикации',
                    'references',
                    'списоклитетатуры',
                    'списокиспользуемойлитературы',
                    'cписокиспользованныхисточников',
                    'ссылкиналитературу']))

#func
def find_toc_fonts_and_sizes(doc):
    xml_str = doc._element.xml
    root = ET.fromstring(xml_str)
    
    # Namespace
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    toc_data = {
        "fonts": [],
        "sizes": []
    }
    
    # Find all sdt elements
    for sdt in root.findall('.//w:sdt', ns):
        # Find all rFonts in this sdt
        for rfont in sdt.findall('.//w:rFonts', ns):
            # Get fonts
            for attr in ['ascii', 'eastAsia', 'hAnsi', 'cs']:
                value = rfont.get(f'{{{ns["w"]}}}{attr}')
                if value:
                    toc_data["fonts"].append(value)
        
        # Find all sz in this sdt
        for sz in sdt.findall('.//w:sz', ns):
            sz_val = sz.get(f'{{{ns["w"]}}}val')
            if sz_val:
                toc_data['sizes'].append(str(sz_val))
        
        # Find all szCs in this sdt
        for szCs in sdt.findall('.//w:szCs', ns):
            szCs_val = szCs.get(f'{{{ns["w"]}}}val')
            if szCs_val:
                toc_data['sizes'].append(str(szCs_val))
    
    toc_data['fonts'] = list(set(toc_data['fonts']))
    toc_data['sizes'] = list(set(toc_data['sizes']))

    return toc_data


def find_paragraphs_followed_by_sdt(doc):
    # Get all paragraphs in document body (python-docx way)
    paragraphs = doc.paragraphs
    
    for i, para in enumerate(paragraphs):
        # Get the XML element for this paragraph
        para_element = para._element
        
        # Get parent and check if next sibling is SDT
        parent = para_element.getparent()
        para_index = parent.index(para_element)
        
        if para_index + 1 < len(parent):
            next_sibling = parent[para_index + 1]
            if etree.QName(next_sibling).localname == 'sdt':
                # Look backwards for first non-empty paragraph
                target_idx = i
                while target_idx >= 0 and not paragraphs[target_idx].text.strip():
                    target_idx -= 1
                
                if target_idx >= 0:
                    return target_idx  # This is the doc.paragraphs index
    
    return 0

def get_gost_citation_by_doi(doi, style="gost"):

    biblio_string = "errorTypeDOITestOccured"

    try:
        citeproc_data = cn.content_negotiation(ids=str(doi), format="bibentry")

        bibtex_file = StringIO(citeproc_data)

        alexlang = "ru"

        if style == "gostnew":

            try:

                match = re.search(r'title\s*=\s*{([^}]+)}', citeproc_data)
                if match:
                    title = match.group(1)
                    alexlang = analyze_script_distribution(title)
            
            except:
                alexlang = "ru"

        if alexlang == "en" and style == "gostnew":
            style = "gostnewen"

        bib_source = BibTeX(bibtex_file)

        BASE_DIR = os.path.dirname(os.path.abspath(__file__))

        #csl_path = os.path.join(BASE_DIR, 'apiv1', 'static', 'csl', f'{style}.csl')
        csl_path = os.getcwd() + '/djangoproject/static/csl/' + str(style) + '.csl'
        
        bib_style = CitationStylesStyle(csl_path, validate=False)
        
        bibliography = CitationStylesBibliography(bib_style, bib_source, formatter.plain)
        
        citation_key = list(bib_source.keys())[0]
        
        citation = Citation([CitationItem(citation_key)])
        bibliography.register(citation)

        bibliography.bibliography()
        
        for item in bibliography.bibliography():
            if hasattr(item, 'text'):
                biblio_string = item.text
            else:
                biblio_string = str(item)
        
        if len(biblio_string) >= 2:
            if biblio_string[:2] == "1.":
                biblio_string = biblio_string[2:]
        if len(biblio_string) >= 3:
            if biblio_string[:3] == "[1]":
                biblio_string = biblio_string[3:]
                
        return biblio_string
        
    except Exception as e:
        print(str(e))
        return biblio_string

def get_first_non_empty_paragraph_index(doc):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            return idx
    return 0

def analyze_script_distribution(text):

    result = "ru"
    
    latin_pattern = re.compile(r'[A-Za-zÀ-ÿ]')  # Includes accented Latin
    cyrillic_pattern = re.compile(r'[\u0400-\u04FF]')  # Cyrillic Unicode block
    
    latin_chars = latin_pattern.findall(text)
    cyrillic_chars = cyrillic_pattern.findall(text)
    
    latin_count = len(latin_chars)
    cyrillic_count = len(cyrillic_chars)
    total_letters = latin_count + cyrillic_count
    
    if total_letters == 0:
        result = "ru"
    elif latin_count > cyrillic_count:
        result = "en"
    elif cyrillic_count > latin_count:
        result = "ru"
    else:
        result = "ru"
    
    return result

def extract_doi_simple(text):
    """
    Extract DOI by finding anything that starts with '10.' and continues
    until the DOI naturally ends (before whitespace, punctuation, or end of string).
    """
    if not text:
        return None
    
    # Pattern explanation:
    # 10\.\d+       - Starts with 10. followed by numbers (prefix)
    # /             - Forward slash separator
    # [^\s<>"']+    - Everything that's not whitespace or common delimiters
    # [^\s<>"'.,;]? - Optional last character that's not punctuation (clean ending)
    
    pattern = r'(10\.\d+/[^\s<>"\']+[^\s<>"\'.,;]?)'
    
    match = re.search(pattern, text)
    if match:
        # Clean up the matched DOI
        doi = match.group(1)
        # Remove any trailing punctuation that might have been included
        doi = re.sub(r'[.,;:]+$', '', doi)
        print(doi)
        return doi
    
    return None

def docx_to_pdf_simple(docx_file, pdf_file=None, format_to="pdf"):
    """
    Simple DOCX to PDF conversion using LibreOffice.
    
    Args:
        docx_file: Path to input .docx file
        pdf_file: Path for output .pdf file (optional)
    
    Returns:
        Path to the created PDF file
    """
    # If no PDF path given, use same name with .pdf extension
    if pdf_file is None:
        if docx_file.endswith(".docx"):
            pdf_file = docx_file[:-5] + ".pdf"
        if docx_file.endswith(".doc"):
            pdf_file = docx_file[:-4] + ".docx"
    
    # Make sure input file exists
    if not os.path.exists(docx_file):
        raise FileNotFoundError(f"File not found: {docx_file}")
    
    # Find LibreOffice on your system
    if sys.platform == 'win32':  # Windows
        # Common Windows paths
        paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        # Check PATH
        soffice = "soffice.exe"
    else:  # macOS/Linux
        # Common macOS/Linux paths
        paths = [
            "/usr/bin/soffice",
            "/usr/local/bin/soffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]
        # Check PATH
        soffice = "soffice"
    
    # Try to find LibreOffice
    libreoffice = None
    for path in paths:
        if os.path.exists(path):
            libreoffice = path
            break
    
    # If not found in common paths, try PATH
    if libreoffice is None:
        try:
            if sys.platform == 'win32':
                result = subprocess.run(['where', soffice], capture_output=True, text=True)
            else:
                result = subprocess.run(['which', soffice], capture_output=True, text=True)
            
            if result.returncode == 0:
                libreoffice = result.stdout.strip()
        except:
            pass
    
    # If still not found, show error
    if libreoffice is None:
        print("LibreOffice not found! Please install it from:")
        print("https://www.libreoffice.org/download/")
        print("\nAfter installing, try again.")
        return None
    
    # Build the command
    output_dir = os.path.dirname(pdf_file) or "."

    if format_to == "pdf":
    
        cmd = [
            libreoffice,
            '--headless',           # Run without showing window
            '--convert-to', 'pdf',  # Convert to PDF
            '--outdir', output_dir, # Where to save PDF
            docx_file              # File to convert
        ]
    
    if format_to == "docx":

        cmd = [
            libreoffice,
            "--headless",
            "--convert-to", "docx",
            "--outdir", output_dir,
            docx_file
        ]


    
    print(f"Converting: {docx_file}")
    print(f"Using: {libreoffice}")
    
    # Run conversion
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0:
            print(f"✓ Created: {pdf_file}")
            return pdf_file
        else:
            print(f"Conversion failed: {result.stderr}")
            return None
            
    except subprocess.TimeoutExpired:
        print("Conversion took too long (over 60 seconds)")
        return None
    except Exception as e:
        print(f"Error: {e}")
        return None

def load_style_lists(style_name="default.xml"):
    """Read style config XML with nested lists and return as dictionary with lists."""
    xml_file = style_name
    
    if not os.path.exists(xml_file):
        raise FileNotFoundError(f"Style config file not found: {xml_file}")
    
    # Parse XML
    tree = etree.parse(xml_file)
    root = tree.getroot()
    
    config = {}
    
    # Get all child elements of StyleConfig
    for element in root:
        # Get tag name
        tag = str(element.tag)
        if '}' in tag:
            tag = tag.split('}')[-1]
        
        # Check if this element has child elements (it's a list wrapper)
        if len(element) > 0:
            # This is a list wrapper like ALLOWED_FONTS or FONT_SIZES
            items = []
            for child in element:
                if child.text:
                    value = child.text.strip()
                    items.append(value)
            config[tag] = items
        else:
            # Regular single value
            if element.text:
                value = element.text.strip()
                config[tag] = value
    
    return config

def add_comments_to_paragraph(doc, para_idx, comment_text, author="Lexify//Style", initials="L.E."):
    try:
        paragraphdoc = doc.paragraphs[para_idx]
        if paragraphdoc.runs:
            comment = doc.add_comment(
                runs=paragraphdoc.runs,
                text=comment_text,
                author=author,
                initials=initials,
            )
    except:
        return False
    return True

def add_comment_to_run(doc, run, comment_text, author="Lexify//Style", initials="L.E."):
    try:
        doc.add_comment(
            runs=[run],
            text=comment_text,
            author=author,
            initials=initials,
        )
        return True
    except Exception as e:
        #print(e)
        return False

def add_comment_to_runs(doc, runs, comment_text, author="Lexify//Style", initials="L.E."):
    try:
        doc.add_comment(
            runs=runs,
            text=comment_text,
            author=author,
            initials=initials,
        )
        return True
    except Exception as e:
        #print(e)
        return False

def insert_emoji_near_boundary(text, emoji="⚠️", from_start=True):
    """
    Insert emoji near the boundary between word parts.
    from_start=True: search left to right, insert before first non-alpha (except hyphen)
    from_start=False: search right to left, insert after first non-alpha (except hyphen)
    """
    if from_start:
        for i, ch in enumerate(text):
            if not (ch.isalpha() or ch == '-'):
                # Insert emoji before this char
                return text[:i] + emoji + text[i:]
        # If no non-alpha found, append emoji at end
        return text + emoji
    else:
        for i in range(len(text) - 1, -1, -1):
            ch = text[i]
            if not (ch.isalpha() or ch == '-'):
                # Insert emoji after this char
                return text[:i+1] + emoji + text[i+1:]
        # If no non-alpha found, prepend emoji at start
        return emoji + text

def find_where_pdf_ends(pdf_path, docx_path, pages=1):
    """
    Find which paragraph in DOCX contains all text from PDF.
    Simple logic: Add paragraphs until cumulative length exceeds PDF length.
    """
    # 1. Get PDF character count (no whitespace)
    pdf_doc = fitz.open(pdf_path)
    pdf_text = ""
    
    for i in range(min(pages, len(pdf_doc))):
        page = pdf_doc.load_page(i)
        pdf_text += page.get_text()
    
    pdf_doc.close()

    print("=========================================")

    print(pdf_text)

    print("=========================================")
    
    # Remove all whitespace
    pdf_chars = len(re.sub(r'\s+', '', pdf_text))
    
    # 2. Add DOCX paragraphs until we exceed PDF length
    doc = Document(docx_path)
    total_chars = 0
    
    for idx, para in enumerate(doc.paragraphs):
        # Remove whitespace from paragraph
        para_text_no_ws = re.sub(r'\s+', '', para.text)
        para_chars = len(para_text_no_ws)
        
        if para_chars == 0:
            continue
        
        # Add to total
        total_chars += para_chars
        
        # Check if we've reached or exceeded PDF length
        if total_chars >= pdf_chars:
            print(para.text)
            print()
            return idx
    
    return len(doc.paragraphs)-1

@dataclass(frozen=True)
class CandidateDOI:
    doi: str
    score: Optional[float]
    title: Optional[str]
    year: Optional[int]
    source: str = "crossref"


def extract_year(text: str) -> Optional[int]:
    m = re.search(r"\b(19|20)\d{2}\b", text)
    return int(m.group(0)) if m else None


def parse_crossref_items(payload: dict[str, Any]) -> list[CandidateDOI]:
    items = payload.get("message", {}).get("items", []) or []
    out: list[CandidateDOI] = []

    for it in items:
        doi = it.get("DOI")
        if not doi:
            continue

        title = None
        titles = it.get("title")
        if isinstance(titles, list) and titles:
            title = titles[0]

        year = None
        issued = it.get("issued", {}).get("date-parts")
        if isinstance(issued, list) and issued and isinstance(issued[0], list) and issued[0]:
            if isinstance(issued[0][0], int):
                year = issued[0][0]

        score = it.get("score")
        out.append(
            CandidateDOI(
                doi=str(doi),
                score=float(score) if score is not None else None,
                title=title,
                year=year,
            )
        )

    return out


def crossref_lookup(citation: str, *, mailto: str, rows: int = 5) -> list[CandidateDOI]:
    params = {
        "query.bibliographic": citation,
        "rows": rows,
        "select": "DOI,title,author,issued,container-title,score",
    }
    headers = {
        # Crossref просит идентифицировать клиента и дать mailto для “polite pool”
        "User-Agent": f"DOIResolverDemo/0.1 (mailto:{mailto})"
    }

    r = requests.get(CROSSREF_WORKS_URL, params=params, headers=headers, timeout=20)
    r.raise_for_status()
    return parse_crossref_items(r.json())


def get_document_properties(file_path, load_style_file="default.xml", target_chars=[], target_words=[], do_spellcheck=True, do_styles=True, do_bibliography=True, bibliography_style="default", skip_paras=-1):

    new_file_path = file_path

    load_style_lists_ = load_style_lists(load_style_file) 

    doc = Document(new_file_path)

    #STATS

    plagiate_string = ''
    char_string = ''

    count_chars = 0
    count_words = 0
    count_sentences = 0

    count_bad_words = 0
    count_bad_chars = 0

    count_bibliography = 0
    count_bad_bibliography = 0
    count_not_doi = 0
    count_suggest_doi = 0
    count_not_right_bibliography = 0

    count_styles_error = 0

    #STATS

    for idx, para in enumerate(doc.paragraphs):
        para_text = para.text
        count_chars += len(para_text)
        count_words += len((re.sub(r'\s+', ' ', para_text).strip()).split(" "))
        count_sentences += sum(1 for char in para_text if char in '.?!')


    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                count_chars += len(cell_text)
                count_words += len((re.sub(r'\s+', ' ', cell_text).strip()).split(" "))
                count_sentences += sum(1 for char in cell_text if char in '.?!')
    
    print('count_chars:')
    print(count_chars)
    print(count_words)
    print(count_sentences)

    #STATS

    if do_spellcheck:

        # Spellcheck
        target_chars = target_chars

        for para_idx, para in enumerate(doc.paragraphs):
            if para_idx <= skip_paras:
                continue
            for run_idx, run in enumerate(para.runs):
                run_elem = run._element
                t_run = run_elem.find(qn('w:t'))
                if t_run is not None:
                    t_run_text = t_run.text

                    if t_run_text is not None:
                    
                        for char in target_chars:
                            if char in t_run_text:
                                run.text = run.text.replace(char, "⚠️" + char + "⚠️")
                                add_comment_to_run(doc, run, f'Допустимо не использовать символ: {char}.', author="Lexify//Spellcheck")
                                count_bad_chars += 1
                                char_string += '<CHARSTRING TYPE="CHAR">' + str(char) + '</CHARSTRING>'
                            

        #Spellcheck
        target_words = target_words

        for para_idx, para in enumerate(doc.paragraphs):
            if para_idx <= skip_paras:
                continue
            for run_idx, run in enumerate(para.runs):
                run_elem = run._element
                t_run = run_elem.find(qn('w:t'))
                if t_run is not None:
                    t_run_text = t_run.text
                    
                    if t_run_text is not None:
                        for word in target_words:
                            if re.search(r'\b' + re.escape(word) + r'\b', t_run_text, re.IGNORECASE):
                            
                                run.text = re.sub(
                                    r'\b' + re.escape(word) + r'\b',
                                    "⚠️" + str(word) + "⚠️",
                                    run.text,
                                    flags=re.IGNORECASE
                                )
                                add_comment_to_run(doc, run, f'Не используйте слово в академическом контексте: {word}.', author="Lexify//Spellcheck")
                                count_bad_words += 1
                                char_string += '<CHARSTRING TYPE="WORD">' + str(word) + '</CHARSTRING>'
                            else:
                                if (run_idx + 1) < len(para.runs):
                                    nextrun = para.runs[run_idx + 1]
                                    nextrun_elem = nextrun._element
                                    t_nextrun = nextrun_elem.find(qn('w:t'))
                                    if t_nextrun is not None:
                                        t_nextrun_text = t_nextrun.text

                                        if t_nextrun_text is not None:

                                            if re.search(r'\b' + re.escape(word) + r'\b', t_nextrun_text, re.IGNORECASE):
                                                continue
                                            else:
                                                if re.search(r'\b' + re.escape(word) + r'\b', t_run_text + t_nextrun_text, re.IGNORECASE):
                                                    run.text = insert_emoji_near_boundary(run.text, emoji="⚠️", from_start=False)
                                                    nextrun.text = insert_emoji_near_boundary(nextrun.text, emoji="⚠️", from_start=True)
                                                    
                                                    add_comment_to_run(doc, run, f'Не используйте слово в академическом контексте: {word}.', author="Lexify//Spellcheck")
                                                    count_bad_words += 1
                                                    char_string += '<CHARSTRING TYPE="WORD">' + str(word) + '</CHARSTRING>'
            
    print(count_bad_chars)
    print(count_bad_words)

    #Bibliography
    if do_bibliography:
        is_reference_start = False
        is_ref_list_idx = []

        for idx in range(len(doc.paragraphs)):
            para = doc.paragraphs[idx]
            para_elem = para._element

            para_text = (para.text).lower()
            para_text = re.sub(r"\s+", "", para_text).strip()
            para_text = re.sub(r"[^a-zа-яё]", "", para_text).strip()

            if para_text in ref_keyword_list:
                is_reference_start = True
                is_ref_list_idx = []
                continue
            
            # Title
            is_it_heading = False
            if para.style.name.startswith('Heading') or para.style.name.startswith('Title'):
                is_it_heading = True
            p_pr = para_elem.find(qn('w:pPr'))
            if p_pr is not None:
                if p_pr.find(qn('w:outlineLvl')) is not None:
                    is_it_heading = True
            
            if is_reference_start == True and is_it_heading == True:
                is_reference_start = False
            
            if is_reference_start == True:
                is_ref_list_idx.append(idx)
        

        for idx in is_ref_list_idx:
            ref_string = doc.paragraphs[idx].text
            if (ref_string).strip() == "":
                continue
            else:
                extract_doi = extract_doi_simple(ref_string)
                count_bibliography += 1
                if extract_doi != None:
                    result_doi = get_gost_citation_by_doi(extract_doi, style=bibliography_style)
                    print("result_doi: ", result_doi)
                    if result_doi != "errorTypeDOITestOccured":
                        if result_doi.strip() != ref_string.strip():
                            add_comments_to_paragraph(doc, idx, f'Проверьте цитирование: {result_doi}', author="Lexify//Bibliography")
                            count_bad_bibliography += 1
                            count_not_right_bibliography += 1
                else:
                    try:
                        mailto = os.environ.get("CROSSREF_MAILTO", "example@example.org")
                        citation = ref_string
                        cands = crossref_lookup(citation, mailto=mailto, rows=1)

                        for c in cands[:1]:
                            print('c.doi')
                            print(c.doi)

                            result_doi = get_gost_citation_by_doi(c.doi, style=bibliography_style)
                            print("result_doi: ", result_doi)
                            if result_doi != "errorTypeDOITestOccured":
                                if result_doi.strip() != ref_string.strip():
                                    add_comments_to_paragraph(doc, idx, f'Возможно: {result_doi}', author="Lexify//Bibliography")
                                    count_bad_bibliography += 1
                                    count_suggest_doi += 1
                    except:
                        add_comments_to_paragraph(doc, idx, f'Вы не указали DOI для библиографической ссылки.', author="Lexify//Bibliography")
                        count_bad_bibliography += 1
                        count_not_doi += 1
    
    print(count_bibliography)
    print(count_bad_bibliography)
    print(count_not_doi)
    print(count_not_right_bibliography)

    #Styles
    if do_styles:
        #Table of content
        tocresult = find_toc_fonts_and_sizes(doc)
        idx_toc = find_paragraphs_followed_by_sdt(doc)
       
        if tocresult['fonts'] != []:
            if not set(tocresult['fonts']).issubset(set(load_style_lists_['TOC_FONTS'])):
               
                std_idx_para = get_first_non_empty_paragraph_index(doc)

                if idx_toc > std_idx_para:
                    std_idx_para = idx_toc
                else:
                    std_idx_para = std_idx_para
                
                print('std_idx_para')
                print(std_idx_para)
                print(doc.paragraphs[std_idx_para].text)
               
                add_comment_to_runs(doc, doc.paragraphs[std_idx_para].runs, f"Шрифт у содержания установлен не верно: {', '.join(list(set(tocresult['fonts'])))}. Используйте: {', '.join(load_style_lists_.get('TOC_FONTS', []))}.")
                count_styles_error += 1
                                    
        if tocresult['sizes'] != []:
            if not set(tocresult['sizes']).issubset(set(load_style_lists_['TOC_SIZE'])):
                std_idx_para = get_first_non_empty_paragraph_index(doc)

                if idx_toc > std_idx_para:
                    std_idx_para = idx_toc
                else:
                    std_idx_para = std_idx_para

                add_comment_to_runs(doc, doc.paragraphs[std_idx_para].runs, f"Размер шрифта у содержания установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(tocresult['sizes'])))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TOC_SIZE', []))}.")
                count_styles_error += 1

        # Tables
        for table in doc.tables:

            # Check table alignment (just once per table)
            tbl_pr = table._element.find(qn('w:tblPr'))
            if tbl_pr is not None:
                jc = tbl_pr.find(qn('w:jc'))
                if jc is not None:
                    table_align = jc.get(qn('w:val'), 'left')  # Default is 'left'
                else:
                    table_align = 'left'  # Default when no jc element exists
                
                if table_align.lower() not in load_style_lists_.get('TABLE_TABLE_ALIGNMENT', []):
                    # Add comment to first cell's first paragraph
                    if table.rows and table.rows[0].cells:
                        first_cell = table.rows[0].cells[0]
                        if first_cell.paragraphs:
                            # f'Используйте: {", ".join(load_style_lists_.get("TABLE_TABLE_ALIGNMENT", []))}.')
                            add_comment_to_runs(doc, first_cell.paragraphs[0].runs,
                                            f'Выравнивание таблицы установлено не верно: {translate_table_options.get(table_align.lower(), table_align.lower())}. '
                                            f'Используйте: {", ".join([translate_table_options.get(align.lower(), align.lower()) for align in load_style_lists_.get("TABLE_TABLE_ALIGNMENT", [])])}.')
                            count_styles_error += 1

            else:
                # No tbl_pr element, default alignment
                if 'left' not in load_style_lists_.get('TABLE_TABLE_ALIGNMENT', []):
                    if table.rows and table.rows[0].cells:
                        first_cell = table.rows[0].cells[0]
                        if first_cell.paragraphs:
                            # f'Используйте: {", ".join(load_style_lists_.get("TABLE_TABLE_ALIGNMENT", []))}.')
                            add_comment_to_runs(doc, first_cell.paragraphs[0].runs,
                                            f'Выравнивание таблицы установлено не верно: слева. '
                                            f'Используйте: {", ".join([translate_table_options.get(align.lower(), align.lower()) for align in load_style_lists_.get("TABLE_TABLE_ALIGNMENT", [])])}.')      
                            count_styles_error += 1

            for row in table.rows:
                for cell in row.cells:
                    
                    # Check vertical alignment for the cell (once per cell, not per paragraph)
                    tc_pr = cell._element.find(qn('w:tcPr'))
                    if tc_pr is not None:
                        v_align_elem = tc_pr.find(qn('w:vAlign'))
                        if v_align_elem is not None:
                            vertical_align = v_align_elem.get(qn('w:val'), 'top')  # Default is 'top' if not specified
                        else:
                            vertical_align = 'top'  # Default when no vAlign element exists
                        
                        if vertical_align.lower() not in load_style_lists_.get('TABLE_VERTICAL_ALIGNMENT', []):
                            # Add comment to first paragraph in cell
                            if cell.paragraphs:
                                # 
                                add_comment_to_runs(doc, cell.paragraphs[0].runs,
                                                f'Вертикальное выравнивание текста у таблицы установлено неверно: {translate_table_options.get(vertical_align.lower(), vertical_align.lower())}. '
                                                f'Используйте: {", ".join([translate_table_options.get(align.lower(), align.lower()) for align in load_style_lists_.get("TABLE_VERTICAL_ALIGNMENT", [])])}.')
                                count_styles_error += 1
                    else:
                        # No tc_pr element, default alignment
                        if 'top' not in load_style_lists_.get('TABLE_VERTICAL_ALIGNMENT', []):
                            if cell.paragraphs:
                                add_comment_to_runs(doc, cell.paragraphs[0].runs, 
                                                f'Вертикальное выравнивание текста у таблицы установлено неверно: сверху. '
                                                f'Используйте: {", ".join([translate_table_options.get(align.lower(), align.lower()) for align in load_style_lists_.get("TABLE_VERTICAL_ALIGNMENT", [])])}.')
                                count_styles_error += 1
                    
                    
                    for idx, para in enumerate(cell.paragraphs):

                        para_text = para.text
                        para_elem = para._element

                        plagiate_string += '<PLAG_PARA_TEXT TYPE="TABLE">' + para_text + '</PLAG_PARA_TEXT>'

                        is_table = True
                    
                        # Empty paragraphs
                        is_para_empty = False
                        if para_text.strip() == "":
                            is_para_empty = True
                    
                        p_pr = para_elem.find(qn('w:pPr'))
                        
                        para_align = None
                        para_color = None
                        para_outline = None
                        para_sz = None
                        para_szCss = None
                        para_i = None
                        para_iCs = None
                        para_strike = None
                        para_u = None
                        para_b = None
                        para_bCs = None
                        para_highlight = None

                        spacing_before = None
                        spacing_after = None
                        spacing_line = None
                        spacing_linerule = None
                        shd_color = None
                        shd_val = None
                        shd_fill = None
                        ind_left = None
                        ind_right = None
                        ind_firstline = None
                        ind_hanging = None
                        font_ascii = None
                        font_hAnsi = None
                        font_eastAsia = None
                        font_cs = None

                        font_list = []
                        fontSize_list = []
                        is_it_italic = False
                        is_it_bold = False

                        if p_pr is not None:
                            jc = p_pr.find(qn('w:jc'))
                            if jc is not None:
                                para_align = jc.get(qn('w:val'))

                                if is_table:
                                    if para_align.lower() not in load_style_lists_['TABLE_ALIGNMENT']:
                                        add_comment_to_runs(doc, para.runs, f"Выравнивание установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}. Используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('TABLE_ALIGNMENT', []))}.")
                                        count_styles_error += 1

                            spacing = p_pr.find(qn('w:spacing'))
                            if spacing is not None:

                                spacing_before = spacing.get(qn('w:before'))
                                if spacing_before is not None:
                                    if is_table:
                                        if spacing_before.lower() not in load_style_lists_['TABLE_SPACINGBEFORE']:
                                            add_comment_to_runs(doc, para.runs, f"Интервал перед текста таблицы установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TABLE_SPACINGBEFORE', []))} пт.")
                                            count_styles_error += 1
                                

                                spacing_after = spacing.get(qn('w:after'))
                                if spacing_after is not None:
                                    if is_table:
                                        if spacing_after.lower() not in load_style_lists_['TABLE_SPACINGAFTER']:
                                            add_comment_to_runs(doc, para.runs, f"Интервал после текста таблицы установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TABLE_SPACINGAFTER', []))} пт.")
                                            count_styles_error += 1

                                spacing_line = spacing.get(qn('w:line'))
                                if spacing_line is not None:
                                    if is_table:
                                        if spacing_line.lower() not in load_style_lists_['TABLE_SPACINGLINE']:
                                            add_comment_to_runs(doc, para.runs, f"Междустрочный интервал текста таблицы установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}. Используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('TABLE_SPACINGLINE', []))}.")
                                            count_styles_error += 1

                                spacing_linerule = spacing.get(qn('w:lineRule'))
                                
                            
                            outline = p_pr.find(qn('w:outlineLvl'))
                            if outline is not None:
                                para_outline = outline.get(qn('w:val'))

                            shd = p_pr.find(qn('w:shd'))
                            if shd is not None:
                                shd_color = shd.get(qn('w:color'))
                                shd_val = shd.get(qn('w:val'))
                                shd_fill = shd.get(qn('w:fill'))
                            

                            ind = p_pr.find(qn('w:ind'))
                            if ind is not None:
                                ind_left = ind.get(qn('w:left'))
                                if ind_left is not None:
                                    if is_table:
                                        if ind_left.lower() not in load_style_lists_['TABLE_INDENTLEFT']:
                                            add_comment_to_runs(doc, para.runs, f"Отступ слева от текста таблицы установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TABLE_INDENTLEFT', []))} см.")
                                            count_styles_error += 1

                                ind_right = ind.get(qn('w:right'))
                                if ind_right is not None:
                                    if is_table:
                                        if ind_right.lower() not in load_style_lists_['TABLE_INDENTRIGHT']:
                                            add_comment_to_runs(doc, para.runs, f"Отступ справа от текста таблицы установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TABLE_INDENTRIGHT', []))} см.")
                                            count_styles_error += 1
                                

                                ind_firstline = ind.get(qn('w:firstLine'))
                                #709 = 1,25
                                if ind_firstline is not None:
                                    if is_table:
                                        if ind_firstline.lower() not in load_style_lists_['TABLE_INDENTFIRSTLINE']:
                                            add_comment_to_runs(doc, para.runs, f"Красная строка у текста таблицы установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TABLE_INDENTFIRSTLINE', []))} см.")
                                            count_styles_error += 1
                                

                                ind_hanging = ind.get(qn('w:hanging'))
                                #709 = 1,25
                                if ind_hanging is not None:
                                    if is_table:
                                        if ind_hanging.lower() not in load_style_lists_['TABLE_INDENTHANGING']:
                                            add_comment_to_runs(doc, para.runs, f"Выступ у текста таблицы установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TABLE_INDENTHANGING', []))} см.")
                                            count_styles_error += 1
                            

                            r_pr = p_pr.find(qn('w:rPr'))
                            if r_pr is not None:

                                r_fonts = r_pr.find(qn('w:rFonts'))
                                if r_fonts is not None:

                                    font_list = []

                                    font_ascii = r_fonts.get(qn('w:ascii'))
                                    if font_ascii is not None:
                                        font_list.append(font_ascii)
                                    font_hAnsi = r_fonts.get(qn('w:hAnsi'))
                                    if font_hAnsi is not None:
                                        font_list.append(font_hAnsi)
                                    font_eastAsia = r_fonts.get(qn('w:eastAsia'))
                                    if font_eastAsia is not None:
                                        font_list.append(font_eastAsia)
                                    font_cs = r_fonts.get(qn('w:cs'))
                                    if font_cs is not None:
                                        font_list.append(font_cs)
                                    

                                    if len(font_list) > 0:
                                        if is_table:
                                            if not set(font_list).issubset(set(load_style_lists_['TABLE_FONTS'])):
                                                add_comment_to_runs(doc, para.runs, f"Шрифт у текста таблицы установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TABLE_FONTS', []))}.")
                                                count_styles_error += 1
                                
                                r_color = r_pr.find(qn('w:color'))
                                if r_color is not None:
                                    para_color = r_color.get(qn('w:val'))
                                    if para_color is not None:
                                        
                                        if is_table:
                                            if para_color not in load_style_lists_['TABLE_COLOR']:
                                                add_comment_to_runs(doc, para.runs, f"Цвет у текста таблицы установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('TABLE_COLOR', []))}.")
                                                count_styles_error += 1

                                fontSize_list = []
                                
                                r_sz = r_pr.find(qn('w:sz'))
                                if r_sz is not None:
                                    para_sz = r_sz.get(qn('w:val'))
                                    if para_sz is not None:
                                        fontSize_list.append(para_sz)
                                
                                r_szCss = r_pr.find(qn('w:szCss'))
                                if r_szCss is not None:
                                    para_szCss = r_szCss.get(qn('w:val'))
                                    if para_szCss is not None:
                                        fontSize_list.append(para_szCss)

                                if len(fontSize_list) > 0:
                                    if is_table:
                                        if not set(fontSize_list).issubset(set(load_style_lists_['TABLE_SIZE'])):
                                            add_comment_to_runs(doc, para.runs, f"Размер шрифта у текста таблицы установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TABLE_SIZE', []))}.")
                                            count_styles_error += 1
                                
                                is_it_italic = False
                                
                                r_i = r_pr.find(qn('w:i'))
                                if r_i is not None:
                                    para_i = True
                                    is_it_italic = True
                            
                                r_iCs = r_pr.find(qn('w:iCs'))
                                if r_iCs is not None:
                                    para_iCs = True
                                    is_it_italic = True
                                
                                if is_it_italic == True:
                                    if is_table:
                                        if "True" not in load_style_lists_['TABLE_ITALIC']:
                                            add_comment_to_runs(doc, para.runs, f'Курсивное начертание у текста таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                                
                                r_strike = r_pr.find(qn('w:strike'))
                                if r_strike is not None:
                                    para_strike = True
                                
                                if para_strike == True:
                                    if is_table:
                                        if "True" not in load_style_lists_['TABLE_STRIKE']:
                                            add_comment_to_runs(doc, para.runs, f'Зачеркнутое начертание у текста таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                        
                                
                                r_u = r_pr.find(qn('w:u'))
                                if r_u is not None:
                                    para_u = True
                                
                                if para_u == True:
                                    if is_table:
                                        if "True" not in load_style_lists_['TABLE_UNDERLINE']:
                                            add_comment_to_runs(doc, para.runs, f'Подчеркнутое начертание у текста таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                                
                                is_it_bold = False

                                r_b = r_pr.find(qn('w:b'))
                                if r_b is not None:
                                    para_b = True
                                    is_it_bold = True
                                
                                r_bCs = r_pr.find(qn('w:bCs'))
                                if r_bCs is not None:
                                    para_bCs = True
                                    is_it_bold = True
                                
                                if is_it_bold == True:
                                    if is_table:
                                        if "True" not in load_style_lists_['TABLE_BOLD']:
                                            add_comment_to_runs(doc, para.runs, f'Полужирное начертание у текста таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                            

                                r_highlight = r_pr.find(qn('w:highlight'))
                                if r_highlight is not None:
                                    para_highlight = True
                                

                                if para_highlight == True:
                                    if is_table:
                                        if "True" not in load_style_lists_['TABLE_HIGHLIGHT']:
                                            add_comment_to_runs(doc, para.runs, f'Выделение цветом у текста таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                                    
                                
                                # Runs

                                for run_idx, run in enumerate(para.runs):
                                    run_elem = run._element

                                    t_run = run_elem.find(qn('w:t'))
                                    if t_run is not None:

                                        is_long_dash = False
                                        if t_run.text is not None and "—" in t_run.text:
                                            is_long_dash = True
                                        
                                        if is_long_dash == True:
                                            if is_table:
                                                if "True" not in load_style_lists_['TABLE_DASH']:
                                                    add_comment_to_run(doc, run, f'Использование длинного тире — у текста таблицы установлено не верно. Не используйте его.')
                                                    count_styles_error += 1

                                    r_pr = run_elem.find(qn('w:rPr'))
                                    if r_pr is not None:

                                        r_fonts = r_pr.find(qn('w:rFonts'))
                                        if r_fonts is not None:

                                            if len(font_list) == 0:
                                        
                                                font_list = []

                                                font_ascii = r_fonts.get(qn('w:ascii'))
                                                if font_ascii is not None:
                                                    font_list.append(font_ascii)
                                                font_hAnsi = r_fonts.get(qn('w:hAnsi'))
                                                if font_hAnsi is not None:
                                                    font_list.append(font_hAnsi)
                                                font_eastAsia = r_fonts.get(qn('w:eastAsia'))
                                                if font_eastAsia is not None:
                                                    font_list.append(font_eastAsia)
                                                font_cs = r_fonts.get(qn('w:cs'))
                                                if font_cs is not None:
                                                    font_list.append(font_cs)
                                                
                                                if len(font_list) > 0:
                                                    if is_table:
                                                        if not set(font_list).issubset(set(load_style_lists_['TABLE_FONTS'])):
                                                            add_comment_to_run(doc, run, f"Шрифт у текста таблицы установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TABLE_FONTS', []))}.")
                                                            count_styles_error += 1
                                    
                                        if len(fontSize_list) == 0:

                                            fontSize_list = []
                                            
                                            r_sz = r_pr.find(qn('w:sz'))
                                            if r_sz is not None:
                                                para_sz = r_sz.get(qn('w:val'))
                                                if para_sz is not None:
                                                    fontSize_list.append(para_sz)
                                            
                                            r_szCss = r_pr.find(qn('w:szCss'))
                                            if r_szCss is not None:
                                                para_szCss = r_szCss.get(qn('w:val'))
                                                if para_szCss is not None:
                                                    fontSize_list.append(para_szCss)

                                            if len(fontSize_list) > 0:
                                                if is_table:
                                                    if not set(fontSize_list).issubset(set(load_style_lists_['TABLE_SIZE'])):
                                                        add_comment_to_run(doc, run, f"Размер шрифта у текста таблицы установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TABLE_SIZE', []))}.")
                                                        count_styles_error += 1
                                    
                                        if para_color is None:
                                            r_color = r_pr.find(qn('w:color'))
                                            if r_color is not None:
                                                para_color = r_color.get(qn('w:val'))
                                                if para_color is not None:
                                                    
                                                    if is_table:
                                                        if para_color not in load_style_lists_['TABLE_COLOR']:
                                                            add_comment_to_run(doc, run, f"Цвет у текста таблицы установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('TABLE_COLOR', []))}.")
                                                            count_styles_error += 1
                                        
                                        if is_it_italic == False:
                                            is_it_italic = False
                                            
                                            r_i = r_pr.find(qn('w:i'))
                                            if r_i is not None:
                                                para_i = True
                                                is_it_italic = True
                                        
                                            r_iCs = r_pr.find(qn('w:iCs'))
                                            if r_iCs is not None:
                                                para_iCs = True
                                                is_it_italic = True
                                            
                                            if is_it_italic == True:
                                                if is_table:
                                                    if "True" not in load_style_lists_['TABLE_ITALIC']:
                                                        add_comment_to_run(doc, run, f'Курсивное начертание у текста таблицы установлено не верно. Не используйте его.')
                                                        count_styles_error += 1
                                    
                                        
                                        if para_strike == None:
                                            r_strike = r_pr.find(qn('w:strike'))
                                            if r_strike is not None:
                                                para_strike = True
                                            
                                            if para_strike == True:
                                                if is_table:
                                                    if "True" not in load_style_lists_['TABLE_STRIKE']:
                                                        add_comment_to_run(doc, run, f'Зачеркнутое начертание у текста таблицы установлено не верно. Не используйте его.')
                                                        count_styles_error += 1
                        
                                            
                                        if para_u == None:
                                            r_u = r_pr.find(qn('w:u'))
                                            if r_u is not None:
                                                para_u = True
                                            
                                            if para_u == True:
                                                if is_table:
                                                    if "True" not in load_style_lists_['TABLE_UNDERLINE']:
                                                        add_comment_to_run(doc, run, f'Подчеркнутое начертание у текста таблицы установлено не верно. Не используйте его.')
                                                        count_styles_error += 1
                                            
                                        if is_it_bold == False:
                                            is_it_bold = False

                                            r_b = r_pr.find(qn('w:b'))
                                            if r_b is not None:
                                                para_b = True
                                                is_it_bold = True
                                            
                                            r_bCs = r_pr.find(qn('w:bCs'))
                                            if r_bCs is not None:
                                                para_bCs = True
                                                is_it_bold = True
                                            
                                            if is_it_bold == True:
                                                if is_table:
                                                    if "True" not in load_style_lists_['TABLE_BOLD']:
                                                        add_comment_to_run(doc, run, f'Полужирное начертание у текста таблицы установлено не верно. Не используйте его.')
                                                        count_styles_error += 1
                                    

                                        if para_highlight == None:
                                            r_highlight = r_pr.find(qn('w:highlight'))
                                            if r_highlight is not None:
                                                para_highlight = True
                                            

                                            if para_highlight == True:
                                                if is_table:
                                                    if "True" not in load_style_lists_['TABLE_HIGHLIGHT']:
                                                        add_comment_to_run(doc, run, f'Выделение цветом у текста таблицы установлено не верно. Не используйте его.')
                                                        count_styles_error += 1
                                        
        # Footers and Headers

        all_header_footer_paras = []

        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                all_header_footer_paras.append(para)

        for section in doc.sections:
            footer = section.footer
            for para in footer.paragraphs:
                all_header_footer_paras.append(para)
        
        for idx, para in enumerate(all_header_footer_paras):
            para_text = para.text
            para_elem = para._element

            plagiate_string += '<PLAG_PARA_TEXT TYPE="HEADER">' + para_text + '</PLAG_PARA_TEXT>'

            is_footerheader = True
        
            p_pr = para_elem.find(qn('w:pPr'))
            
        
            para_sz = None
            para_szCss = None
        
            font_ascii = None
            font_hAnsi = None
            font_eastAsia = None
            font_cs = None

            font_list = []
            fontSize_list = []
        

            if p_pr is not None:
                r_pr = p_pr.find(qn('w:rPr'))
                if r_pr is not None:

                    r_fonts = r_pr.find(qn('w:rFonts'))
                    if r_fonts is not None:

                        font_list = []

                        font_ascii = r_fonts.get(qn('w:ascii'))
                        if font_ascii is not None:
                            font_list.append(font_ascii)
                        font_hAnsi = r_fonts.get(qn('w:hAnsi'))
                        if font_hAnsi is not None:
                            font_list.append(font_hAnsi)
                        font_eastAsia = r_fonts.get(qn('w:eastAsia'))
                        if font_eastAsia is not None:
                            font_list.append(font_eastAsia)
                        font_cs = r_fonts.get(qn('w:cs'))
                        if font_cs is not None:
                            font_list.append(font_cs)

                
                        

                        if len(font_list) > 0:
                            if is_footerheader:
                                if not set(font_list).issubset(set(load_style_lists_['TABLE_FONTS'])):
                                    add_comment_to_runs(doc, doc.paragraphs[get_first_non_empty_paragraph_index(doc)].runs, f"Шрифт у footer установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TABLE_FONTS', []))}.")
                                    count_styles_error += 1
                    
                
                    fontSize_list = []
                    
                    r_sz = r_pr.find(qn('w:sz'))
                    if r_sz is not None:
                        para_sz = r_sz.get(qn('w:val'))
                        if para_sz is not None:
                            fontSize_list.append(para_sz)
                    
                    r_szCss = r_pr.find(qn('w:szCss'))
                    if r_szCss is not None:
                        para_szCss = r_szCss.get(qn('w:val'))
                        if para_szCss is not None:
                            fontSize_list.append(para_szCss)
                    
                    

                    if len(fontSize_list) > 0:
                        if is_footerheader:
                            if not set(fontSize_list).issubset(set(load_style_lists_['TABLE_SIZE'])):
                                add_comment_to_runs(doc, doc.paragraphs[get_first_non_empty_paragraph_index(doc)].runs, f"Размер шрифта у footer установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TABLE_SIZE', []))}.")
                                count_styles_error += 1
                    
                
                    for run_idx, run in enumerate(para.runs):
                        run_elem = run._element
                        
                        r_pr = run_elem.find(qn('w:rPr'))
                        if r_pr is not None:

                            r_fonts = r_pr.find(qn('w:rFonts'))
                            if r_fonts is not None:

                                if 1==1:
                            
                                    font_list = []

                                    font_ascii = r_fonts.get(qn('w:ascii'))
                                    if font_ascii is not None:
                                        font_list.append(font_ascii)
                                    font_hAnsi = r_fonts.get(qn('w:hAnsi'))
                                    if font_hAnsi is not None:
                                        font_list.append(font_hAnsi)
                                    font_eastAsia = r_fonts.get(qn('w:eastAsia'))
                                    if font_eastAsia is not None:
                                        font_list.append(font_eastAsia)
                                    font_cs = r_fonts.get(qn('w:cs'))
                                    if font_cs is not None:
                                        font_list.append(font_cs)
                                    
                                    if len(font_list) > 0:
                                        if is_footerheader:
                                            if not set(font_list).issubset(set(load_style_lists_['TABLE_FONTS'])):
                                                add_comment_to_runs(doc, doc.paragraphs[get_first_non_empty_paragraph_index(doc)].runs, f"Шрифт у footer установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TABLE_FONTS', []))}.")
                                                count_styles_error += 1
                        
                            if 1==1:

                                fontSize_list = []
                                
                                r_sz = r_pr.find(qn('w:sz'))
                                if r_sz is not None:
                                    para_sz = r_sz.get(qn('w:val'))
                                    if para_sz is not None:
                                        fontSize_list.append(para_sz)
                                
                                r_szCss = r_pr.find(qn('w:szCss'))
                                if r_szCss is not None:
                                    para_szCss = r_szCss.get(qn('w:val'))
                                    if para_szCss is not None:
                                        fontSize_list.append(para_szCss)
                                
                        

                                if len(fontSize_list) > 0:
                                    if is_footerheader:
                                        if not set(fontSize_list).issubset(set(load_style_lists_['TABLE_SIZE'])):
                                            add_comment_to_runs(doc, doc.paragraphs[get_first_non_empty_paragraph_index(doc)].runs, f"Размер шрифта у footer установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TABLE_SIZE', []))}.")
                                            count_styles_error += 1





        #Page Settings
        sections = doc.sections
        for section in sections:

            section_buffer_list = []

            page_height = section.page_height
            if page_height:
                if not any(abs(float(page_height) - float(h)) <= float(h) * 0.01 for h in load_style_lists_['PAGE_HEIGHT']):
                    #add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), f"В документе неверная высота страницы: {float(page_height) * 2.54 / 914400:.2f} см. Используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_HEIGHT', []))} см.")
                    section_buffer_list.append(f"В документе неверная высота страницы: {float(page_height) * 2.54 / 914400:.2f} см., используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_HEIGHT', []))} см.")
                    count_styles_error += 1

            page_width = section.page_width
            if page_width:
                if not any(abs(float(page_width) - float(h)) <= float(h) * 0.01 for h in load_style_lists_['PAGE_WIDTH']):
                    #add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), f"В документе неверная ширина страницы: {float(page_width) * 2.54 / 914400:.2f} см. Используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_WIDTH', []))} см.")
                    section_buffer_list.append(f"В документе неверная ширина страницы: {float(page_width) * 2.54 / 914400:.2f} см., используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_WIDTH', []))} см.")
                    count_styles_error += 1

            left_margin = section.left_margin
            if left_margin:
                if not any(abs(float(left_margin) - float(h)) <= float(h) * 0.01 for h in load_style_lists_['PAGE_MARGINLEFT']):
                    #add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), f"В документе неверный отступ слева на странице: {float(left_margin) * 2.54 / 914400:.2f} см. Используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_MARGINLEFT', []))} см.")
                    section_buffer_list.append(f"В документе неверный отступ слева на странице: {float(left_margin) * 2.54 / 914400:.2f} см., используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_MARGINLEFT', []))} см.")
                    count_styles_error += 1

            right_margin = section.right_margin
            if right_margin:
                if not any(abs(float(right_margin) - float(h)) <= float(h) * 0.01 for h in load_style_lists_['PAGE_MARGINRIGHT']):
                    #add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), f"В документе неверный отступ справа на странице: {float(right_margin) * 2.54 / 914400:.2f} см. Используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_MARGINRIGHT', []))} см.")
                    section_buffer_list.append(f"В документе неверный отступ справа на странице: {float(right_margin) * 2.54 / 914400:.2f} см., используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_MARGINRIGHT', []))} см.")
                    count_styles_error += 1

            top_margin = section.top_margin
            if top_margin:
                if not any(abs(float(top_margin) - float(h)) <= float(h) * 0.01 for h in load_style_lists_['PAGE_MARGINTOP']):
                    #add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), f"В документе неверный отступ сверху на странице: {float(top_margin) * 2.54 / 914400:.2f} см. Используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_MARGINTOP', []))} см.")
                    section_buffer_list.append(f"В документе неверный отступ сверху на странице: {float(top_margin) * 2.54 / 914400:.2f} см., используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_MARGINTOP', []))} см.")
                    count_styles_error += 1

            bottom_margin = section.bottom_margin
            if bottom_margin:
                if not any(abs(float(bottom_margin) - float(h)) <= float(h) * 0.01 for h in load_style_lists_['PAGE_MARGINBOTTOM']):
                    #add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), f"В документе неверный отступ снизу на странице: {float(bottom_margin) * 2.54 / 914400:.2f} см. Используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_MARGINBOTTOM', []))} см.")
                    section_buffer_list.append(f"В документе неверный отступ снизу на странице: {float(bottom_margin) * 2.54 / 914400:.2f} см., используйте: {', '.join(str(round(float(a)/ 914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_MARGINBOTTOM', []))} см.")
                    count_styles_error += 1

            header_distance = section.header_distance
            if header_distance:
                if not any(abs(float(header_distance) - float(h)) <= float(h) * 0.01 for h in load_style_lists_['PAGE_HEADERDISTANCE']):
                    #add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), f"В документе неверная высота верхнего колонтитула: {float(header_distance) * 2.54 / 914400:.2f} см. Используйте: {', '.join(str(round(float(a)/914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_HEADERDISTANCE', []))} см.")
                    section_buffer_list.append(f"В документе неверная высота верхнего колонтитула: {float(header_distance) * 2.54 / 914400:.2f} см., используйте: {', '.join(str(round(float(a)/914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_HEADERDISTANCE', []))} см.")
                    count_styles_error += 1

            footer_distance = section.footer_distance
            if footer_distance:
                if not any(abs(float(footer_distance) - float(h)) <= float(h) * 0.01 for h in load_style_lists_['PAGE_FOOTERDISTANCE']):
                    #add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), f"В документе неверная высота нижнего колонтитула: {float(footer_distance) * 2.54 / 914400:.2f} см. Используйте: {', '.join(str(round(float(a)/914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_FOOTERDISTANCE', []))} см.")
                    section_buffer_list.append(f"В документе неверная высота нижнего колонтитула: {float(footer_distance) * 2.54 / 914400:.2f} см., используйте: {', '.join(str(round(float(a)/914400 * 2.54, 2)) for a in load_style_lists_.get('PAGE_FOOTERDISTANCE', []))} см.")
                    count_styles_error += 1

            if len(section_buffer_list) > 0:
                add_comments_to_paragraph(doc, get_first_non_empty_paragraph_index(doc), " - " + "\n - ".join(section_buffer_list))    

        
        # Reading doc styles       

        prev_has_drawing = False
        
        for idx, para in enumerate(doc.paragraphs):
            if idx <= skip_paras:
                continue

            list_buffer_notes = []

            para_text = para.text
            para_elem = para._element

            plagiate_string += '<PLAG_PARA_TEXT TYPE="TEXT">' + para_text + '</PLAG_PARA_TEXT>'
        
            # Empty paragraphs
            is_para_empty = False
            if para_text.strip() == "":
                is_para_empty = True

            # Drawings
            has_drawing = bool(para_elem.findall('.//' + qn('w:drawing')))

            if not has_drawing:
                if is_para_empty:
                    if 'False' in load_style_lists_['EMPTY_PARAGRAPH']:
                        #add_comments_to_paragraph(doc, idx, f'Не используйте пустые абзацы.')
                        list_buffer_notes.append(f'Не используйте пустые абзацы.')
                        count_styles_error += 1
            
            # Drawing captions
            is_drawing_caption = False
            if prev_has_drawing:
                is_drawing_caption = True
            prev_has_drawing = has_drawing

            # Table heading
            next_elem = para_elem.getnext()
            is_table_heading = False
            if next_elem is not None and next_elem.tag.endswith('}tbl'):
                is_table_heading = True

            # Title
            is_it_heading = False
            if para.style.name.startswith('Heading') or para.style.name.startswith('Title'):
                is_it_heading = True
            p_pr = para_elem.find(qn('w:pPr'))
            if p_pr is not None:
                if p_pr.find(qn('w:outlineLvl')) is not None:
                    is_it_heading = True
            

            para_style_name = para.style.name
        
            p_pr = para_elem.find(qn('w:pPr'))
            
            para_align = None
            para_color = None
            para_outline = None
            para_sz = None
            para_szCss = None
            para_i = None
            para_iCs = None
            para_strike = None
            para_u = None
            para_b = None
            para_bCs = None
            para_highlight = None

            spacing_before = None
            spacing_after = None
            spacing_line = None
            spacing_linerule = None
            shd_color = None
            shd_val = None
            shd_fill = None
            ind_left = None
            ind_right = None
            ind_firstline = None
            ind_hanging = None
            font_ascii = None
            font_hAnsi = None
            font_eastAsia = None
            font_cs = None

            font_list = []
            fontSize_list = []
            is_it_italic = False
            is_it_bold = False

            if p_pr is not None:
                jc = p_pr.find(qn('w:jc'))
                if jc is not None:
                    para_align = jc.get(qn('w:val'))

                    if has_drawing:
                        if para_align.lower() not in load_style_lists_['DRAWING_ALIGNMENT']:
                            #add_comments_to_paragraph(doc, idx, f"Выравнивание рисунка установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}. Используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('DRAWING_ALIGNMENT', []))}.")
                            list_buffer_notes.append(f"Выравнивание рисунка установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}., используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('DRAWING_ALIGNMENT', []))}.")
                            count_styles_error += 1

                    elif is_drawing_caption:
                        if para_align.lower() not in load_style_lists_['DRAWINGCAPTION_ALIGNMENT']:
                            #add_comments_to_paragraph(doc, idx, f"Выравнивание подписи к рисунку установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}. Используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('DRAWINGCAPTION_ALIGNMENT', []))}.")
                            list_buffer_notes.append(f"Выравнивание подписи к рисунку установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}., используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('DRAWINGCAPTION_ALIGNMENT', []))}.")
                            count_styles_error += 1

                    elif is_table_heading:
                        if para_align.lower() not in load_style_lists_['TABLEHEADING_ALIGNMENT']:
                            #add_comments_to_paragraph(doc, idx, f"Выравнивание заголовка таблицы установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}. Используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('TABLEHEADING_ALIGNMENT', []))}.")
                            list_buffer_notes.append(f"Выравнивание заголовка таблицы установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}., используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('TABLEHEADING_ALIGNMENT', []))}.")
                            count_styles_error += 1

                    elif is_it_heading:
                        if para_align.lower() not in load_style_lists_['HEADING_ALIGNMENT']:
                            #add_comments_to_paragraph(doc, idx, f"Выравнивание заголовка установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}. Используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('HEADING_ALIGNMENT', []))}.")
                            list_buffer_notes.append(f"Выравнивание заголовка установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}., используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('HEADING_ALIGNMENT', []))}.")
                            count_styles_error += 1
                    else:
                        if para_align.lower() not in load_style_lists_['TEXT_ALIGNMENT']:
                            #add_comments_to_paragraph(doc, idx, f"Выравнивание текста установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}. Используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('TEXT_ALIGNMENT', []))}.")
                            list_buffer_notes.append(f"Выравнивание текста установлено не верно: {translate_alignmnet_options.get(para_align.lower(), para_align)}., используйте: {', '.join(translate_alignmnet_options.get(align.lower(), align) for align in load_style_lists_.get('TEXT_ALIGNMENT', []))}.")
                            count_styles_error += 1
                    
                
                
                spacing = p_pr.find(qn('w:spacing'))
                if spacing is not None:

                    spacing_before = spacing.get(qn('w:before'))
                    if spacing_before is not None:
                        if has_drawing:
                            if spacing_before.lower() not in load_style_lists_['DRAWING_SPACINGBEFORE']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал перед рисунком установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('DRAWING_SPACINGBEFORE', []))} пт.")
                                list_buffer_notes.append(f"Интервал перед рисунком установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('DRAWING_SPACINGBEFORE', []))} пт.")
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if spacing_before.lower() not in load_style_lists_['DRAWINGCAPTION_SPACINGBEFORE']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал перед подписью к рисунку установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_SPACINGBEFORE', []))} пт.")
                                list_buffer_notes.append(f"Интервал перед подписью к рисунку установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_SPACINGBEFORE', []))} пт.")
                                count_styles_error += 1
                        elif is_table_heading:
                            if spacing_before.lower() not in load_style_lists_['TABLEHEADING_SPACINGBEFORE']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал перед заголовком таблицы установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TABLEHEADING_SPACINGBEFORE', []))} пт.")
                                list_buffer_notes.append(f"Интервал перед заголовком таблицы установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TABLEHEADING_SPACINGBEFORE', []))} пт.")
                                count_styles_error += 1

                        elif is_it_heading:
                            if spacing_before.lower() not in load_style_lists_['HEADING_SPACINGBEFORE']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал перед заголовком установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('HEADING_SPACINGBEFORE', []))} пт.")
                                list_buffer_notes.append(f"Интервал перед заголовком установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('HEADING_SPACINGBEFORE', []))} пт.")
                                count_styles_error += 1
                        else:
                            if spacing_before.lower() not in load_style_lists_['TEXT_SPACINGBEFORE']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал перед текстом установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TEXT_SPACINGBEFORE', []))} пт.")
                                list_buffer_notes.append(f"Интервал перед текстом установлен не верно: {round(float(spacing_before.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TEXT_SPACINGBEFORE', []))} пт.")
                                count_styles_error += 1

                    spacing_after = spacing.get(qn('w:after'))
                    if spacing_after is not None:
                        if has_drawing:
                            if spacing_after.lower() not in load_style_lists_['DRAWING_SPACINGAFTER']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал после рисунка установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('DRAWING_SPACINGAFTER', []))} пт.")
                                list_buffer_notes.append(f"Интервал после рисунка установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('DRAWING_SPACINGAFTER', []))} пт.")
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if spacing_after.lower() not in load_style_lists_['DRAWINGCAPTION_SPACINGAFTER']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал после подписи к рисунку установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_SPACINGAFTER', []))} пт.")
                                list_buffer_notes.append(f"Интервал после подписи к рисунку установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_SPACINGAFTER', []))} пт.")
                                count_styles_error += 1

                        elif is_table_heading:
                            if spacing_after.lower() not in load_style_lists_['TABLEHEADING_SPACINGAFTER']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал после заголовка таблицы установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TABLEHEADING_SPACINGAFTER', []))} пт.")
                                list_buffer_notes.append(f"Интервал после заголовка таблицы установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TABLEHEADING_SPACINGAFTER', []))} пт.")
                                count_styles_error += 1

                        elif is_it_heading:
                            if spacing_after.lower() not in load_style_lists_['HEADING_SPACINGAFTER']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал после заголовка установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('HEADING_SPACINGAFTER', []))} пт.")
                                list_buffer_notes.append(f"Интервал после заголовка установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('HEADING_SPACINGAFTER', []))} пт.")
                                count_styles_error += 1
                        else:
                            if spacing_after.lower() not in load_style_lists_['TEXT_SPACINGAFTER']:
                                #add_comments_to_paragraph(doc, idx, f"Интервал после текста установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт. Используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TEXT_SPACINGAFTER', []))} пт.")
                                list_buffer_notes.append(f"Интервал после текста установлен не верно: {round(float(spacing_after.lower())*12/240,2):.2f} пт., используйте: {', '.join(f'{round(float(align.lower())*12/240,2):.2f}' for align in load_style_lists_.get('TEXT_SPACINGAFTER', []))} пт.")
                                count_styles_error += 1

                    spacing_line = spacing.get(qn('w:line'))
                    if spacing_line is not None:
                        if has_drawing:
                            if spacing_line.lower() not in load_style_lists_['DRAWING_SPACINGLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Междустрочный интервал рисунка установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}. Используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('DRAWING_SPACINGLINE', []))}.")
                                list_buffer_notes.append(f"Междустрочный интервал рисунка установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}., используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('DRAWING_SPACINGLINE', []))}.")
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if spacing_line.lower() not in load_style_lists_['DRAWINGCAPTION_SPACINGLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Междустрочный интервал подписи к рисунку установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}. Используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_SPACINGLINE', []))}.")
                                list_buffer_notes.append(f"Междустрочный интервал подписи к рисунку установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}., используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_SPACINGLINE', []))}.")
                                count_styles_error += 1

                        elif is_table_heading:
                            if spacing_line.lower() not in load_style_lists_['TABLEHEADING_SPACINGLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Междустрочный интервал заголовка таблицы установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}. Используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('TABLEHEADING_SPACINGLINE', []))}.")
                                list_buffer_notes.append(f"Междустрочный интервал заголовка таблицы установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}., используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('TABLEHEADING_SPACINGLINE', []))}.")
                                count_styles_error += 1

                        elif is_it_heading:
                            if spacing_line.lower() not in load_style_lists_['HEADING_SPACINGLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Междустрочный интервал заголовка установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}. Используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('HEADING_SPACINGLINE', []))}.")
                                list_buffer_notes.append(f"Междустрочный интервал заголовка установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}., используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('HEADING_SPACINGLINE', []))}.")
                                count_styles_error += 1

                        else:
                            if spacing_line.lower() not in load_style_lists_['TEXT_SPACINGLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Междустрочный интервал текста установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}. Используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('TEXT_SPACINGLINE', []))}.")
                                list_buffer_notes.append(f"Междустрочный интервал текста установлен не верно: {round(float(spacing_line.lower())*1/240,2):.2f}., используйте: {', '.join(f'{(round(float(align.lower())*1/240,2)):.2f}' for align in load_style_lists_.get('TEXT_SPACINGLINE', []))}.")
                                count_styles_error += 1

                    spacing_linerule = spacing.get(qn('w:lineRule'))
                    
                
                outline = p_pr.find(qn('w:outlineLvl'))
                if outline is not None:
                    para_outline = outline.get(qn('w:val'))

                shd = p_pr.find(qn('w:shd'))
                if shd is not None:
                    shd_color = shd.get(qn('w:color'))
                    shd_val = shd.get(qn('w:val'))
                    shd_fill = shd.get(qn('w:fill'))
                

                ind = p_pr.find(qn('w:ind'))
                if ind is not None:
                    ind_left = ind.get(qn('w:left'))
                    if ind_left is not None:
                        if has_drawing:
                            if ind_left.lower() not in load_style_lists_['DRAWING_INDENTLEFT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ слева от рисунка установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('DRAWING_INDENTLEFT', []))} см.")
                                list_buffer_notes.append(f"Отступ слева от рисунка установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('DRAWING_INDENTLEFT', []))} см.")
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if ind_left.lower() not in load_style_lists_['DRAWINGCAPTION_INDENTLEFT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ слева от подписи к рисунку установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_INDENTLEFT', []))} см.")
                                list_buffer_notes.append(f"Отступ слева от подписи к рисунку установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_INDENTLEFT', []))} см.")
                                count_styles_error += 1

                        elif is_table_heading:
                            if ind_left.lower() not in load_style_lists_['TABLEHEADING_INDENTLEFT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ слева от заголовка таблицы установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TABLEHEADING_INDENTLEFT', []))} см.")
                                list_buffer_notes.append(f"Отступ слева от заголовка таблицы установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TABLEHEADING_INDENTLEFT', []))} см.")
                                count_styles_error += 1

                        elif is_it_heading:
                            if ind_left.lower() not in load_style_lists_['HEADING_INDENTLEFT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ слева от заголовка установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('HEADING_INDENTLEFT', []))} см.")
                                list_buffer_notes.append(f"Отступ слева от заголовка установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('HEADING_INDENTLEFT', []))} см.")
                                count_styles_error += 1

                        else:
                            if ind_left.lower() not in load_style_lists_['TEXT_INDENTLEFT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ слева от текста установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TEXT_INDENTLEFT', []))} см.")
                                list_buffer_notes.append(f"Отступ слева от текста установлен не верно: {(float(ind_left.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TEXT_INDENTLEFT', []))} см.")
                                count_styles_error += 1


                    ind_right = ind.get(qn('w:right'))
                    if ind_right is not None:
                        if has_drawing:
                            if ind_right.lower() not in load_style_lists_['DRAWING_INDENTRIGHT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ справа от рисунка установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('DRAWING_INDENTRIGHT', []))} см.")
                                list_buffer_notes.append(f"Отступ справа от рисунка установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('DRAWING_INDENTRIGHT', []))} см.")
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if ind_right.lower() not in load_style_lists_['DRAWINGCAPTION_INDENTRIGHT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ справа от подписи к рисунку установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_INDENTRIGHT', []))} см.")
                                list_buffer_notes.append(f"Отступ справа от подписи к рисунку установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_INDENTRIGHT', []))} см.")
                                count_styles_error += 1

                        elif is_table_heading:
                            if ind_right.lower() not in load_style_lists_['TABLEHEADING_INDENTRIGHT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ справа от заголовка таблицы установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TABLEHEADING_INDENTRIGHT', []))} см.")
                                list_buffer_notes.append(f"Отступ справа от заголовка таблицы установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TABLEHEADING_INDENTRIGHT', []))} см.")
                                count_styles_error += 1

                        elif is_it_heading:
                            if ind_right.lower() not in load_style_lists_['HEADING_INDENTRIGHT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ справа от заголовка установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('HEADING_INDENTRIGHT', []))} см.")
                                list_buffer_notes.append(f"Отступ справа от заголовка установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('HEADING_INDENTRIGHT', []))} см.")
                                count_styles_error += 1

                        else:
                            if ind_right.lower() not in load_style_lists_['TEXT_INDENTRIGHT']:
                                #add_comments_to_paragraph(doc, idx, f"Отступ справа от текста установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TEXT_INDENTRIGHT', []))} см.")
                                list_buffer_notes.append(f"Отступ справа от текста установлен не верно: {(float(ind_right.lower())*2/1134):.2f} см., используйте: {', '.join(f'{(float(align.lower())*2/1134):.2f}' for align in load_style_lists_.get('TEXT_INDENTRIGHT', []))} см.")
                                count_styles_error += 1


                    ind_firstline = ind.get(qn('w:firstLine'))
                    #709 = 1,25
                    if ind_firstline is not None:
                        if has_drawing:
                            if ind_firstline.lower() not in load_style_lists_['DRAWING_INDENTFIRSTLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Красная строка у рисунка установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('DRAWING_INDENTFIRSTLINE', []))} см.")
                                list_buffer_notes.append(f"Красная строка у рисунка установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('DRAWING_INDENTFIRSTLINE', []))} см.")
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if ind_firstline.lower() not in load_style_lists_['DRAWINGCAPTION_INDENTFIRSTLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Красная строка у подписи к рисунку установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_INDENTFIRSTLINE', []))} см.")
                                list_buffer_notes.append(f"Красная строка у подписи к рисунку установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_INDENTFIRSTLINE', []))} см.")
                                count_styles_error += 1

                        elif is_table_heading:
                            if ind_firstline.lower() not in load_style_lists_['TABLEHEADING_INDENTFIRSTLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Красная строка у заголовка таблицы установлен не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TABLEHEADING_INDENTFIRSTLINE', []))} см.")
                                list_buffer_notes.append(f"Красная строка у заголовка таблицы установлен не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TABLEHEADING_INDENTFIRSTLINE', []))} см.")
                                count_styles_error += 1

                        elif is_it_heading:
                            if ind_firstline.lower() not in load_style_lists_['HEADING_INDENTFIRSTLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Красная строка у заголовка установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('HEADING_INDENTFIRSTLINE', []))} см.")
                                list_buffer_notes.append(f"Красная строка у заголовка установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('HEADING_INDENTFIRSTLINE', []))} см.")
                                count_styles_error += 1

                        else:
                            if ind_firstline.lower() not in load_style_lists_['TEXT_INDENTFIRSTLINE']:
                                #add_comments_to_paragraph(doc, idx, f"Красная строка у текста установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TEXT_INDENTFIRSTLINE', []))} см.")
                                list_buffer_notes.append(f"Красная строка у текста установлена не верно: {(float(ind_firstline.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TEXT_INDENTFIRSTLINE', []))} см.")
                                count_styles_error += 1
                                #resulttest = highlight_runs(doc.paragraphs[idx].runs, "yellow")
                                #print(resulttest)

                    ind_hanging = ind.get(qn('w:hanging'))
                    #709 = 1,25
                    if ind_hanging is not None:
                        if has_drawing:
                            if ind_hanging.lower() not in load_style_lists_['DRAWING_INDENTHANGING']:
                                #add_comments_to_paragraph(doc, idx, f"Выступ у рисунка установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('DRAWING_INDENTHANGING', []))} см.")
                                list_buffer_notes.append(f"Выступ у рисунка установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('DRAWING_INDENTHANGING', []))} см.")
                                count_styles_error += 1
                        elif is_drawing_caption:
                            if ind_hanging.lower() not in load_style_lists_['DRAWINGCAPTION_INDENTHANGING']:
                                #add_comments_to_paragraph(doc, idx, f"Выступ у подписи к рисунку установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_INDENTHANGING', []))} см.")
                                list_buffer_notes.append(f"Выступ у подписи к рисунку установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('DRAWINGCAPTION_INDENTHANGING', []))} см.")
                                count_styles_error += 1
                        elif is_table_heading:
                            if ind_hanging.lower() not in load_style_lists_['TABLEHEADING_INDENTHANGING']:
                                #add_comments_to_paragraph(doc, idx, f"Выступ у заголовка таблицы установлен не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TABLEHEADING_INDENTHANGING', []))} см.")
                                list_buffer_notes.append(f"Выступ у заголовка таблицы установлен не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TABLEHEADING_INDENTHANGING', []))} см.")
                                count_styles_error += 1
                        elif is_it_heading:
                            if ind_hanging.lower() not in load_style_lists_['HEADING_INDENTHANGING']:
                                #add_comments_to_paragraph(doc, idx, f"Выступ у заголовка установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('HEADING_INDENTHANGING', []))} см.")
                                list_buffer_notes.append(f"Выступ у заголовка установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('HEADING_INDENTHANGING', []))} см.")
                                count_styles_error += 1
                        else:
                            if ind_hanging.lower() not in load_style_lists_['TEXT_INDENTHANGING']:
                                #add_comments_to_paragraph(doc, idx, f"Выступ у текста установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см. Используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TEXT_INDENTHANGING', []))} см.")
                                list_buffer_notes.append(f"Выступ у текста установлена не верно: {(float(ind_hanging.lower())*1.25/709):.2f} см., используйте: {', '.join(f'{(float(align.lower())*1.25/709):.2f}' for align in load_style_lists_.get('TEXT_INDENTHANGING', []))} см.")
                                count_styles_error += 1

                

                r_pr = p_pr.find(qn('w:rPr'))
                if r_pr is not None:

                    r_fonts = r_pr.find(qn('w:rFonts'))
                    if r_fonts is not None:

                        font_list = []

                        font_ascii = r_fonts.get(qn('w:ascii'))
                        if font_ascii is not None:
                            font_list.append(font_ascii)
                        font_hAnsi = r_fonts.get(qn('w:hAnsi'))
                        if font_hAnsi is not None:
                            font_list.append(font_hAnsi)
                        font_eastAsia = r_fonts.get(qn('w:eastAsia'))
                        if font_eastAsia is not None:
                            font_list.append(font_eastAsia)
                        font_cs = r_fonts.get(qn('w:cs'))
                        if font_cs is not None:
                            font_list.append(font_cs)
                        

                        if len(font_list) > 0:
                            if has_drawing:
                                if not set(font_list).issubset(set(load_style_lists_['DRAWING_FONTS'])):
                                    #add_comments_to_paragraph(doc, idx, f"Шрифт у рисунка установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('DRAWING_FONTS', []))}.")
                                    list_buffer_notes.append(f"Шрифт у рисунка установлен не верно: {', '.join(list(set(font_list)))}., используйте: {', '.join(load_style_lists_.get('DRAWING_FONTS', []))}.")
                                    count_styles_error += 1

                            elif is_drawing_caption:
                                if not set(font_list).issubset(set(load_style_lists_['DRAWINGCAPTION_FONTS'])):
                                    #add_comments_to_paragraph(doc, idx, f"Шрифт у подписи к рисунку установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('DRAWINGCAPTION_FONTS', []))}.")
                                    list_buffer_notes.append(f"Шрифт у подписи к рисунку установлен не верно: {', '.join(list(set(font_list)))}., используйте: {', '.join(load_style_lists_.get('DRAWINGCAPTION_FONTS', []))}.")
                                    count_styles_error += 1

                            elif is_table_heading:
                                if not set(font_list).issubset(set(load_style_lists_['TABLEHEADING_FONTS'])):
                                    #add_comments_to_paragraph(doc, idx, f"Шрифт у заголовка таблицы установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TABLEHEADING_FONTS', []))}.")
                                    list_buffer_notes.append(f"Шрифт у заголовка таблицы установлен не верно: {', '.join(list(set(font_list)))}., используйте: {', '.join(load_style_lists_.get('TABLEHEADING_FONTS', []))}.")
                                    count_styles_error += 1

                            elif is_it_heading:
                                if not set(font_list).issubset(set(load_style_lists_['HEADING_FONTS'])):
                                    #add_comments_to_paragraph(doc, idx, f"Шрифт у заголовка установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('HEADING_FONTS', []))}.")
                                    list_buffer_notes.append(f"Шрифт у заголовка установлен не верно: {', '.join(list(set(font_list)))}., используйте: {', '.join(load_style_lists_.get('HEADING_FONTS', []))}.")
                                    count_styles_error += 1

                            else:
                                if not set(font_list).issubset(set(load_style_lists_['TEXT_FONTS'])):
                                    #add_comments_to_paragraph(doc, idx, f"Шрифт у текста установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TEXT_FONTS', []))}.")
                                    list_buffer_notes.append(f"Шрифт у текста установлен не верно: {', '.join(list(set(font_list)))}., используйте: {', '.join(load_style_lists_.get('TEXT_FONTS', []))}.")
                                    count_styles_error += 1
                                    #resulttest = highlight_runs(doc.paragraphs[idx].runs, "yellow")
                                    #print(resulttest)
                    
                    r_color = r_pr.find(qn('w:color'))
                    if r_color is not None:
                        para_color = r_color.get(qn('w:val'))
                        if para_color is not None:
                            
                            if has_drawing:
                                if para_color not in load_style_lists_['DRAWING_COLOR']:
                                    #add_comments_to_paragraph(doc, idx, f"Цвет у рисунка установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('DRAWING_COLOR', []))}.")
                                    list_buffer_notes.append(f"Цвет у рисунка установлен не верно: #{', #'.join([para_color])}., используйте: #{', #'.join(load_style_lists_.get('DRAWING_COLOR', []))}.")
                                    count_styles_error += 1

                            elif is_drawing_caption:
                                if para_color not in load_style_lists_['DRAWINGCAPTION_COLOR']:
                                    #add_comments_to_paragraph(doc, idx, f"Цвет у подписи к рисунку установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('DRAWINGCAPTION_COLOR', []))}.")
                                    list_buffer_notes.append(f"Цвет у подписи к рисунку установлен не верно: #{', #'.join([para_color])}., используйте: #{', #'.join(load_style_lists_.get('DRAWINGCAPTION_COLOR', []))}.")
                                    count_styles_error += 1

                            elif is_table_heading:
                                if para_color not in load_style_lists_['TABLEHEADING_COLOR']:
                                    #add_comments_to_paragraph(doc, idx, f"Цвет у заголовка таблицы установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('TABLEHEADING_COLOR', []))}.")
                                    list_buffer_notes.append(f"Цвет у заголовка таблицы установлен не верно: #{', #'.join([para_color])}., используйте: #{', #'.join(load_style_lists_.get('TABLEHEADING_COLOR', []))}.")
                                    count_styles_error += 1

                            elif is_it_heading:
                                if para_color not in load_style_lists_['HEADING_COLOR']:
                                    #add_comments_to_paragraph(doc, idx, f"Цвет у заголовка установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('HEADING_COLOR', []))}.")
                                    list_buffer_notes.append(f"Цвет у заголовка установлен не верно: #{', #'.join([para_color])}., используйте: #{', #'.join(load_style_lists_.get('HEADING_COLOR', []))}.")
                                    count_styles_error += 1

                            else:
                                if para_color not in load_style_lists_['TEXT_COLOR']:
                                    #add_comments_to_paragraph(doc, idx, f"Цвет у текста установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('TEXT_COLOR', []))}.")
                                    list_buffer_notes.append(f"Цвет у текста установлен не верно: #{', #'.join([para_color])}., используйте: #{', #'.join(load_style_lists_.get('TEXT_COLOR', []))}.")
                                    count_styles_error += 1


                    fontSize_list = []
                    
                    r_sz = r_pr.find(qn('w:sz'))
                    if r_sz is not None:
                        para_sz = r_sz.get(qn('w:val'))
                        if para_sz is not None:
                            fontSize_list.append(para_sz)
                    
                    r_szCss = r_pr.find(qn('w:szCss'))
                    if r_szCss is not None:
                        para_szCss = r_szCss.get(qn('w:val'))
                        if para_szCss is not None:
                            fontSize_list.append(para_szCss)

                    if len(fontSize_list) > 0:
                        if has_drawing:
                            if not set(fontSize_list).issubset(set(load_style_lists_['DRAWING_SIZE'])):
                                #dd_comments_to_paragraph(doc, idx, f"Размер шрифта у рисунка установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('DRAWING_SIZE', []))}.")
                                list_buffer_notes.append(f"Размер шрифта у рисунка установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}., используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('DRAWING_SIZE', []))}.")
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if not set(fontSize_list).issubset(set(load_style_lists_['DRAWINGCAPTION_SIZE'])):
                                #add_comments_to_paragraph(doc, idx, f"Размер шрифта у подписи к рисунку установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('DRAWINGCAPTION_SIZE', []))}.")
                                list_buffer_notes.append(f"Размер шрифта у подписи к рисунку установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}., используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('DRAWINGCAPTION_SIZE', []))}.")
                                count_styles_error += 1

                        elif is_table_heading:
                            if not set(fontSize_list).issubset(set(load_style_lists_['TABLEHEADING_SIZE'])):
                                #add_comments_to_paragraph(doc, idx, f"Размер шрифта у заголовка таблицы установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TABLEHEADING_SIZE', []))}.")
                                list_buffer_notes.append(f"Размер шрифта у заголовка таблицы установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}., используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TABLEHEADING_SIZE', []))}.")
                                count_styles_error += 1

                        elif is_it_heading:
                            if not set(fontSize_list).issubset(set(load_style_lists_['HEADING_SIZE'])):
                                #add_comments_to_paragraph(doc, idx, f"Размер шрифта у заголовка установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('HEADING_SIZE', []))}.")
                                list_buffer_notes.append(f"Размер шрифта у заголовка установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}., используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('HEADING_SIZE', []))}.")
                                count_styles_error += 1
                        else:
                            if not set(fontSize_list).issubset(set(load_style_lists_['TEXT_SIZE'])):
                                #add_comments_to_paragraph(doc, idx, f"Размер шрифта у текста установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TEXT_SIZE', []))}.")
                                list_buffer_notes.append(f"Размер шрифта у текста установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}., используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TEXT_SIZE', []))}.")
                                count_styles_error += 1

                    
                    is_it_italic = False
                    
                    r_i = r_pr.find(qn('w:i'))
                    if r_i is not None:
                        para_i = True
                        is_it_italic = True
                
                    r_iCs = r_pr.find(qn('w:iCs'))
                    if r_iCs is not None:
                        para_iCs = True
                        is_it_italic = True
                    
                    if is_it_italic == True:
                        if has_drawing:
                            if "True" not in load_style_lists_['DRAWING_ITALIC']:
                                #add_comments_to_paragraph(doc, idx, f'Курсивное начертание у рисунка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Курсивное начертание у рисунка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if "True" not in load_style_lists_['DRAWINGCAPTION_ITALIC']:
                                #add_comments_to_paragraph(doc, idx, f'Курсивное начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Курсивное начертание у подписи к рисунку установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_table_heading:
                            if "True" not in load_style_lists_['TABLEHEADING_ITALIC']:
                                #add_comments_to_paragraph(doc, idx, f'Курсивное начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Курсивное начертание у заголовка таблицы установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_it_heading:
                            if "True" not in load_style_lists_['HEADING_ITALIC']:
                                #add_comments_to_paragraph(doc, idx, f'Курсивное начертание у заголовка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Курсивное начертание у заголовка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        else:
                            if "True" not in load_style_lists_['TEXT_ITALIC']:
                                #add_comments_to_paragraph(doc, idx, f'Курсивное начертание у текста установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Курсивное начертание у текста установлено не верно, не используйте его.')
                                count_styles_error += 1
                    
                    
                    r_strike = r_pr.find(qn('w:strike'))
                    if r_strike is not None:
                        para_strike = True
                    
                    if para_strike == True:
                        if has_drawing:
                            if "True" not in load_style_lists_['DRAWING_STRIKE']:
                                #add_comments_to_paragraph(doc, idx, f'Зачеркнутое начертание у рисунка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Зачеркнутое начертание у рисунка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if "True" not in load_style_lists_['DRAWINGCAPTION_STRIKE']:
                                #add_comments_to_paragraph(doc, idx, f'Зачеркнутое начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Зачеркнутое начертание у подписи к рисунку установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_table_heading:
                            if "True" not in load_style_lists_['TABLEHEADING_STRIKE']:
                                #add_comments_to_paragraph(doc, idx, f'Зачеркнутое начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Зачеркнутое начертание у заголовка таблицы установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_it_heading:
                            if "True" not in load_style_lists_['HEADING_STRIKE']:
                                #add_comments_to_paragraph(doc, idx, f'Зачеркнутое начертание у заголовка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Зачеркнутое начертание у заголовка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        else:
                            if "True" not in load_style_lists_['TEXT_STRIKE']:
                                #add_comments_to_paragraph(doc, idx, f'Зачеркнутое начертание у текста установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Зачеркнутое начертание у текста установлено не верно, не используйте его.')
                                count_styles_error += 1
                    
            
                    
                    r_u = r_pr.find(qn('w:u'))
                    if r_u is not None:
                        para_u = True
                    
                    if para_u == True:
                        if has_drawing:
                            if "True" not in load_style_lists_['DRAWING_UNDERLINE']:
                                #add_comments_to_paragraph(doc, idx, f'Подчеркнутое начертание у рисунка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Подчеркнутое начертание у рисунка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if "True" not in load_style_lists_['DRAWINGCAPTION_UNDERLINE']:
                                #add_comments_to_paragraph(doc, idx, f'Подчеркнутое начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Подчеркнутое начертание у подписи к рисунку установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_table_heading:
                            if "True" not in load_style_lists_['TABLEHEADING_UNDERLINE']:
                                #add_comments_to_paragraph(doc, idx, f'Подчеркнутое начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Подчеркнутое начертание у заголовка таблицы установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_it_heading:
                            if "True" not in load_style_lists_['HEADING_UNDERLINE']:
                                #add_comments_to_paragraph(doc, idx, f'Подчеркнутое начертание у заголовка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Подчеркнутое начертание у заголовка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        else:
                            if "True" not in load_style_lists_['TEXT_UNDERLINE']:
                                #add_comments_to_paragraph(doc, idx, f'Подчеркнутое начертание у текста установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Подчеркнутое начертание у текста установлено не верно, не используйте его.')
                                count_styles_error += 1
                    
                    
                    is_it_bold = False

                    r_b = r_pr.find(qn('w:b'))
                    if r_b is not None:
                        para_b = True
                        is_it_bold = True
                    
                    r_bCs = r_pr.find(qn('w:bCs'))
                    if r_bCs is not None:
                        para_bCs = True
                        is_it_bold = True
                    
                    if is_it_bold == True:
                        if has_drawing:
                            if "True" not in load_style_lists_['DRAWING_BOLD']:
                                #add_comments_to_paragraph(doc, idx, f'Полужирное начертание у рисунка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Полужирное начертание у рисунка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if "True" not in load_style_lists_['DRAWINGCAPTION_BOLD']:
                                #add_comments_to_paragraph(doc, idx, f'Полужирное начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Полужирное начертание у подписи к рисунку установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_table_heading:
                            if "True" not in load_style_lists_['TABLEHEADING_BOLD']:
                                #add_comments_to_paragraph(doc, idx, f'Полужирное начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Полужирное начертание у заголовка таблицы установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_it_heading:
                            if "True" not in load_style_lists_['HEADING_BOLD']:
                                #add_comments_to_paragraph(doc, idx, f'Полужирное начертание у заголовка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Полужирное начертание у заголовка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        else:
                            if "True" not in load_style_lists_['TEXT_BOLD']:
                                #add_comments_to_paragraph(doc, idx, f'Полужирное начертание у текста установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Полужирное начертание у текста установлено не верно, не используйте его.')
                                count_styles_error += 1
                    
                

                    r_highlight = r_pr.find(qn('w:highlight'))
                    if r_highlight is not None:
                        para_highlight = True
                    

                    if para_highlight == True:
                        if has_drawing:
                            if "True" not in load_style_lists_['DRAWING_HIGHLIGHT']:
                                #add_comments_to_paragraph(doc, idx, f'Выделение цветом у рисунка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Выделение цветом у рисунка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_drawing_caption:
                            if "True" not in load_style_lists_['DRAWINGCAPTION_HIGHLIGHT']:
                                #add_comments_to_paragraph(doc, idx, f'Выделение цветом у подписи к рисунку установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Выделение цветом у подписи к рисунку установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_table_heading:
                            if "True" not in load_style_lists_['TABLEHEADING_HIGHLIGHT']:
                                #add_comments_to_paragraph(doc, idx, f'Выделение цветом у заголовка таблицы установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Выделение цветом у заголовка таблицы установлено не верно, не используйте его.')
                                count_styles_error += 1

                        elif is_it_heading:
                            if "True" not in load_style_lists_['HEADING_HIGHLIGHT']:
                                #add_comments_to_paragraph(doc, idx, f'Выделение цветом у заголовка установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Выделение цветом у заголовка установлено не верно, не используйте его.')
                                count_styles_error += 1

                        else:
                            if "True" not in load_style_lists_['TEXT_HIGHLIGHT']:
                                #add_comments_to_paragraph(doc, idx, f'Выделение цветом у текста установлено не верно. Не используйте его.')
                                list_buffer_notes.append(f'Выделение цветом у текста установлено не верно, не используйте его.')
                                count_styles_error += 1
                    

            
                    
                    # Runs

                    for run_idx, run in enumerate(para.runs):
                        run_elem = run._element

                        run_buffer_list = []

                        t_run = run_elem.find(qn('w:t'))
                        if t_run is not None:

                            is_long_dash = False
                            if t_run.text is not None and "—" in t_run.text:
                                is_long_dash = True
                            
                            if is_long_dash == True:
                                if has_drawing:
                                    if "True" not in load_style_lists_['DRAWING_DASH']:
                                        #add_comment_to_run(doc, run, f'Использование длинного тире — у рисунка установлено не верно. Не используйте его.')
                                        run_buffer_list.append(f'Использование длинного тире — у рисунка установлено не верно. Не используйте его.')
                                        count_styles_error += 1

                                elif is_drawing_caption:
                                    if "True" not in load_style_lists_['DRAWINGCAPTION_DASH']:
                                        #add_comment_to_run(doc, run, f'Использование длинного тире — у подписи к рисунку установлено не верно. Не используйте его.')
                                        run_buffer_list.append(f'Использование длинного тире — у подписи к рисунку установлено не верно. Не используйте его.')
                                        count_styles_error += 1

                                elif is_table_heading:
                                    if "True" not in load_style_lists_['TABLEHEADING_DASH']:
                                        #add_comment_to_run(doc, run, f'Использование длинного тире — у заголовка таблицы установлено не верно. Не используйте его.')
                                        run_buffer_list.append(f'Использование длинного тире — у заголовка таблицы установлено не верно. Не используйте его.')
                                        count_styles_error += 1

                                elif is_it_heading:
                                    if "True" not in load_style_lists_['HEADING_DASH']:
                                        #add_comment_to_run(doc, run, f'Использование длинного тире — у заголовка установлено не верно. Не используйте его.')
                                        run_buffer_list.append(f'Использование длинного тире — у заголовка установлено не верно. Не используйте его.')
                                        count_styles_error += 1

                                else:
                                    if "True" not in load_style_lists_['TEXT_DASH']:
                                        #add_comment_to_run(doc, run, f'Использование длинного тире — у текста установлено не верно. Не используйте его.')
                                        run_buffer_list.append(f'Использование длинного тире — у текста установлено не верно. Не используйте его.')
                                        count_styles_error += 1
                            

                        r_pr = run_elem.find(qn('w:rPr'))
                        if r_pr is not None:

                            r_fonts = r_pr.find(qn('w:rFonts'))
                            if r_fonts is not None:

                                if len(font_list) == 0:
                            
                                    font_list = []

                                    font_ascii = r_fonts.get(qn('w:ascii'))
                                    if font_ascii is not None:
                                        font_list.append(font_ascii)
                                    font_hAnsi = r_fonts.get(qn('w:hAnsi'))
                                    if font_hAnsi is not None:
                                        font_list.append(font_hAnsi)
                                    font_eastAsia = r_fonts.get(qn('w:eastAsia'))
                                    if font_eastAsia is not None:
                                        font_list.append(font_eastAsia)
                                    font_cs = r_fonts.get(qn('w:cs'))
                                    if font_cs is not None:
                                        font_list.append(font_cs)
                                    
                                    if len(font_list) > 0:
                                        if has_drawing:
                                            if not set(font_list).issubset(set(load_style_lists_['DRAWING_FONTS'])):
                                                #add_comment_to_run(doc, run, f"Шрифт у рисунка установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('DRAWING_FONTS', []))}.")
                                                run_buffer_list.append(f"Шрифт у рисунка установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('DRAWING_FONTS', []))}.")
                                                count_styles_error += 1

                                        elif is_drawing_caption:
                                            if not set(font_list).issubset(set(load_style_lists_['DRAWINGCAPTION_FONTS'])):
                                                #add_comment_to_run(doc, run, f"Шрифт у подписи к рисунку установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('DRAWINGCAPTION_FONTS', []))}.")
                                                run_buffer_list.append(f"Шрифт у подписи к рисунку установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('DRAWINGCAPTION_FONTS', []))}.")
                                                count_styles_error += 1

                                        elif is_table_heading:
                                            if not set(font_list).issubset(set(load_style_lists_['TABLEHEADING_FONTS'])):
                                                #add_comment_to_run(doc, run, f"Шрифт у заголовка таблицы установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TABLEHEADING_FONTS', []))}.")
                                                run_buffer_list.append(f"Шрифт у заголовка таблицы установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TABLEHEADING_FONTS', []))}.")
                                                count_styles_error += 1

                                        elif is_it_heading:
                                            if not set(font_list).issubset(set(load_style_lists_['HEADING_FONTS'])):
                                                #add_comment_to_run(doc, run, f"Шрифт у заголовка установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('HEADING_FONTS', []))}.")
                                                run_buffer_list.append(f"Шрифт у заголовка установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('HEADING_FONTS', []))}.")
                                                count_styles_error += 1

                                        else:
                                            if not set(font_list).issubset(set(load_style_lists_['TEXT_FONTS'])):
                                                #add_comment_to_run(doc, run, f"Шрифт у текста установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TEXT_FONTS', []))}.")
                                                run_buffer_list.append(f"Шрифт у текста установлен не верно: {', '.join(list(set(font_list)))}. Используйте: {', '.join(load_style_lists_.get('TEXT_FONTS', []))}.")
                                                count_styles_error += 1
                            
                        
                            if len(fontSize_list) == 0:

                                fontSize_list = []
                                
                                r_sz = r_pr.find(qn('w:sz'))
                                if r_sz is not None:
                                    para_sz = r_sz.get(qn('w:val'))
                                    if para_sz is not None:
                                        fontSize_list.append(para_sz)
                                
                                r_szCss = r_pr.find(qn('w:szCss'))
                                if r_szCss is not None:
                                    para_szCss = r_szCss.get(qn('w:val'))
                                    if para_szCss is not None:
                                        fontSize_list.append(para_szCss)

                                if len(fontSize_list) > 0:
                                    if has_drawing:
                                        if not set(fontSize_list).issubset(set(load_style_lists_['DRAWING_SIZE'])):
                                            #add_comment_to_run(doc, run, f"Размер шрифта у рисунка установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('DRAWING_SIZE', []))}.")
                                            run_buffer_list.append(f"Размер шрифта у рисунка установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('DRAWING_SIZE', []))}.")
                                            count_styles_error += 1

                                    elif is_drawing_caption:
                                        if not set(fontSize_list).issubset(set(load_style_lists_['DRAWINGCAPTION_SIZE'])):
                                            #add_comment_to_run(doc, run, f"Размер шрифта у подписи к рисунку установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('DRAWINGCAPTION_SIZE', []))}.")
                                            run_buffer_list.append(f"Размер шрифта у подписи к рисунку установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('DRAWINGCAPTION_SIZE', []))}.")
                                            count_styles_error += 1

                                    elif is_table_heading:
                                        if not set(fontSize_list).issubset(set(load_style_lists_['TABLEHEADING_SIZE'])):
                                            #add_comment_to_run(doc, run, f"Размер шрифта у заголовка таблицы установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TABLEHEADING_SIZE', []))}.")
                                            run_buffer_list.append(f"Размер шрифта у заголовка таблицы установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TABLEHEADING_SIZE', []))}.")
                                            count_styles_error += 1

                                    elif is_it_heading:
                                        if not set(fontSize_list).issubset(set(load_style_lists_['HEADING_SIZE'])):
                                            #add_comment_to_run(doc, run, f"Размер шрифта у заголовка установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('HEADING_SIZE', []))}.")
                                            run_buffer_list.append(f"Размер шрифта у заголовка установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('HEADING_SIZE', []))}.")
                                            count_styles_error += 1

                                    else:
                                        if not set(fontSize_list).issubset(set(load_style_lists_['TEXT_SIZE'])):
                                            #add_comment_to_run(doc, run, f"Размер шрифта у текста установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TEXT_SIZE', []))}.")
                                            run_buffer_list.append(f"Размер шрифта у текста установлен не верно: {', '.join(str(float(v) / 2) for v in list(set(fontSize_list)))}. Используйте: {', '.join(str(float(v) / 2) for v in load_style_lists_.get('TEXT_SIZE', []))}.")
                                            count_styles_error += 1

                        
                            if para_color is None:
                                r_color = r_pr.find(qn('w:color'))
                                if r_color is not None:
                                    para_color = r_color.get(qn('w:val'))
                                    if para_color is not None:
                                        
                                        if has_drawing:
                                            if para_color not in load_style_lists_['DRAWING_COLOR']:
                                                #add_comment_to_run(doc, run, f"Цвет у рисунка установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('DRAWING_COLOR', []))}.")
                                                run_buffer_list.append(f"Цвет у рисунка установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('DRAWING_COLOR', []))}.")
                                                count_styles_error += 1

                                        elif is_drawing_caption:
                                            if para_color not in load_style_lists_['DRAWINGCAPTION_COLOR']:
                                                #add_comment_to_run(doc, run, f"Цвет у подписи к рисунку установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('DRAWINGCAPTION_COLOR', []))}.")
                                                run_buffer_list.append(f"Цвет у подписи к рисунку установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('DRAWINGCAPTION_COLOR', []))}.")
                                                count_styles_error += 1

                                        elif is_table_heading:
                                            if para_color not in load_style_lists_['TABLEHEADING_COLOR']:
                                                #add_comment_to_run(doc, run, f"Цвет у заголовка таблицы установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('TABLEHEADING_COLOR', []))}.")
                                                run_buffer_list.append(f"Цвет у заголовка таблицы установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('TABLEHEADING_COLOR', []))}.")
                                                count_styles_error += 1

                                        elif is_it_heading:
                                            if para_color not in load_style_lists_['HEADING_COLOR']:
                                                #add_comment_to_run(doc, run, f"Цвет у заголовка установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('HEADING_COLOR', []))}.")
                                                run_buffer_list.append(f"Цвет у заголовка установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('HEADING_COLOR', []))}.")
                                                count_styles_error += 1

                                        else:
                                            if para_color not in load_style_lists_['TEXT_COLOR']:
                                                #add_comment_to_run(doc, run, f"Цвет у текста установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('TEXT_COLOR', []))}.")
                                                run_buffer_list.append(f"Цвет у текста установлен не верно: #{', #'.join([para_color])}. Используйте: #{', #'.join(load_style_lists_.get('TEXT_COLOR', []))}.")
                                                count_styles_error += 1

                            
                            if is_it_italic == False:
                                is_it_italic = False
                                
                                r_i = r_pr.find(qn('w:i'))
                                if r_i is not None:
                                    para_i = True
                                    is_it_italic = True
                            
                                r_iCs = r_pr.find(qn('w:iCs'))
                                if r_iCs is not None:
                                    para_iCs = True
                                    is_it_italic = True
                                
                                if is_it_italic == True:
                                    if has_drawing:
                                        if "True" not in load_style_lists_['DRAWING_ITALIC']:
                                            #add_comment_to_run(doc, run, f'Курсивное начертание у рисунка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Курсивное начертание у рисунка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_drawing_caption:
                                        if "True" not in load_style_lists_['DRAWINGCAPTION_ITALIC']:
                                            #add_comment_to_run(doc, run, f'Курсивное начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Курсивное начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_table_heading:
                                        if "True" not in load_style_lists_['TABLEHEADING_ITALIC']:
                                            #add_comment_to_run(doc, run, f'Курсивное начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Курсивное начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_it_heading:
                                        if "True" not in load_style_lists_['HEADING_ITALIC']:
                                            #add_comment_to_run(doc, run, f'Курсивное начертание у заголовка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Курсивное начертание у заголовка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    else:
                                        if "True" not in load_style_lists_['TEXT_ITALIC']:
                                            #add_comment_to_run(doc, run, f'Курсивное начертание у текста установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Курсивное начертание у текста установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                                
                        
                            
                            if para_strike == None:
                                r_strike = r_pr.find(qn('w:strike'))
                                if r_strike is not None:
                                    para_strike = True
                                
                                if para_strike == True:
                                    if has_drawing:
                                        if "True" not in load_style_lists_['DRAWING_STRIKE']:
                                            #add_comment_to_run(doc, run, f'Зачеркнутое начертание у рисунка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Зачеркнутое начертание у рисунка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_drawing_caption:
                                        if "True" not in load_style_lists_['DRAWINGCAPTION_STRIKE']:
                                            #add_comment_to_run(doc, run, f'Зачеркнутое начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Зачеркнутое начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_table_heading:
                                        if "True" not in load_style_lists_['TABLEHEADING_STRIKE']:
                                            #add_comment_to_run(doc, run, f'Зачеркнутое начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Зачеркнутое начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_it_heading:
                                        if "True" not in load_style_lists_['HEADING_STRIKE']:
                                            #add_comment_to_run(doc, run, f'Зачеркнутое начертание у заголовка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Зачеркнутое начертание у заголовка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    else:
                                        if "True" not in load_style_lists_['TEXT_STRIKE']:
                                            #add_comment_to_run(doc, run, f'Зачеркнутое начертание у текста установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Зачеркнутое начертание у текста установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                                
            
                                
                            if para_u == None:
                                r_u = r_pr.find(qn('w:u'))
                                if r_u is not None:
                                    para_u = True
                                
                                if para_u == True:
                                    if has_drawing:
                                        if "True" not in load_style_lists_['DRAWING_UNDERLINE']:
                                            #add_comment_to_run(doc, run, f'Подчеркнутое начертание у рисунка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Подчеркнутое начертание у рисунка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_drawing_caption:
                                        if "True" not in load_style_lists_['DRAWINGCAPTION_UNDERLINE']:
                                            #add_comment_to_run(doc, run, f'Подчеркнутое начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Подчеркнутое начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_table_heading:
                                        if "True" not in load_style_lists_['TABLEHEADING_UNDERLINE']:
                                            #add_comment_to_run(doc, run, f'Подчеркнутое начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Подчеркнутое начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_it_heading:
                                        if "True" not in load_style_lists_['HEADING_UNDERLINE']:
                                            #add_comment_to_run(doc, run, f'Подчеркнутое начертание у заголовка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Подчеркнутое начертание у заголовка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    else:
                                        if "True" not in load_style_lists_['TEXT_UNDERLINE']:
                                            #add_comment_to_run(doc, run, f'Подчеркнутое начертание у текста установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Подчеркнутое начертание у текста установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                                
                                
                            if is_it_bold == False:
                                is_it_bold = False

                                r_b = r_pr.find(qn('w:b'))
                                if r_b is not None:
                                    para_b = True
                                    is_it_bold = True
                                
                                r_bCs = r_pr.find(qn('w:bCs'))
                                if r_bCs is not None:
                                    para_bCs = True
                                    is_it_bold = True
                                
                                if is_it_bold == True:
                                    if has_drawing:
                                        if "True" not in load_style_lists_['DRAWING_BOLD']:
                                            #add_comment_to_run(doc, run, f'Полужирное начертание у рисунка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Полужирное начертание у рисунка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_drawing_caption:
                                        if "True" not in load_style_lists_['DRAWINGCAPTION_BOLD']:
                                            #add_comment_to_run(doc, run, f'Полужирное начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Полужирное начертание у подписи к рисунку установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_table_heading:
                                        if "True" not in load_style_lists_['TABLEHEADING_BOLD']:
                                            #add_comment_to_run(doc, run, f'Полужирное начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Полужирное начертание у заголовка таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_it_heading:
                                        if "True" not in load_style_lists_['HEADING_BOLD']:
                                            #add_comment_to_run(doc, run, f'Полужирное начертание у заголовка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Полужирное начертание у заголовка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    else:
                                        if "True" not in load_style_lists_['TEXT_BOLD']:
                                            #add_comment_to_run(doc, run, f'Полужирное начертание у текста установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Полужирное начертание у текста установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                                
                        

                            if para_highlight == None:
                                r_highlight = r_pr.find(qn('w:highlight'))
                                if r_highlight is not None:
                                    para_highlight = True
                                

                                if para_highlight == True:
                                    if has_drawing:
                                        if "True" not in load_style_lists_['DRAWING_HIGHLIGHT']:
                                            #add_comment_to_run(doc, run, f'Выделение цветом у рисунка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Выделение цветом у рисунка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_drawing_caption:
                                        if "True" not in load_style_lists_['DRAWINGCAPTION_HIGHLIGHT']:
                                            #add_comment_to_run(doc, run, f'Выделение цветом у подписи к рисунку установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Выделение цветом у подписи к рисунку установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_table_heading:
                                        if "True" not in load_style_lists_['TABLEHEADING_HIGHLIGHT']:
                                            #add_comment_to_run(doc, run, f'Выделение цветом у заголовка таблицы установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Выделение цветом у заголовка таблицы установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    elif is_it_heading:
                                        if "True" not in load_style_lists_['HEADING_HIGHLIGHT']:
                                            #add_comment_to_run(doc, run, f'Выделение цветом у заголовка установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Выделение цветом у заголовка установлено не верно. Не используйте его.')
                                            count_styles_error += 1

                                    else:
                                        if "True" not in load_style_lists_['TEXT_HIGHLIGHT']:
                                            #add_comment_to_run(doc, run, f'Выделение цветом у текста установлено не верно. Не используйте его.')
                                            run_buffer_list.append(f'Выделение цветом у текста установлено не верно. Не используйте его.')
                                            count_styles_error += 1
                        
                        
                        
                        if len(run_buffer_list) > 0:
                            add_comments_to_paragraph(doc, run, " - " + "\n - ".join(run_buffer_list))  


            if len(list_buffer_notes) > 0:
                add_comments_to_paragraph(doc, idx, " - " + "\n - ".join(list_buffer_notes))    


    # also check contents (oglavlenie)

    output_path = new_file_path + ".withNotes.docx"

    if new_file_path.endswith(".docx"):
        output_path = new_file_path[:-5] + ".withNotes.docx"

    doc.save(output_path)

    return output_path, plagiate_string, count_chars, count_words, count_sentences, count_bad_words, count_bad_chars, count_bibliography, count_bad_bibliography, count_not_doi, count_not_right_bibliography, count_styles_error, char_string, count_suggest_doi


def get_pdf_pages_text(pdf_path: str) -> dict:
    pages_text = {}
    
    pdf_document = fitz.open(pdf_path)
    
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        text = page.get_text()
        text = re.sub(r'\s+', ' ', text)
        
        pages_text[page_num + 1] = text
    
    pdf_document.close()
    
    return pages_text



def get_docx_paragraphs_text(docx_path: str) -> dict:
    doc = Document(docx_path)
    paragraphs_dict = {}
    
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        paragraphs_dict[idx] = text
    
    return paragraphs_dict



def map_paragraphs_to_pages_sequential(pdf_pages: dict, docx_paragraphs: dict) -> dict:
   
    def clean_text(text):
        if not isinstance(text, str):
            text = str(text)
        # Normalize whitespace and special characters
        text = text.replace('\xa0', ' ').replace('\t', ' ')
        text = ' '.join(text.split())
        return text.strip()
    
    # Clean all texts
    pdf_clean = {page: clean_text(text) for page, text in pdf_pages.items()}
    docx_clean = {idx: clean_text(text) for idx, text in docx_paragraphs.items()}
    
    # Sort paragraph indices (they should already be in order from 0, 1, 2...)
    paragraph_indices = sorted(docx_clean.keys())
    
    # Initialize result
    result = {}
    current_page = 1
    last_match_position = -1  # Track position in text for sequential matching
    
    for para_idx in paragraph_indices:
        para_text = docx_clean[para_idx]
        
        # Skip empty or very short paragraphs
        if not para_text or len(para_text) < 5:
            continue
            
        best_page = None
        
        # Try to find this paragraph in pages starting from current page
        for page_num in range(current_page, max(pdf_clean.keys()) + 1):
            if page_num not in pdf_clean:
                continue
                
            page_text = pdf_clean[page_num]
            
            # Check if paragraph is in this page's text
            if para_text in page_text:
                best_page = page_num
                break
            
            # For longer paragraphs, check if at least 80% of words match
            if len(para_text) > 30:
                para_words = set(para_text.lower().split())
                page_words = set(page_text.lower().split())
                common_words = para_words.intersection(page_words)
                
                if len(para_words) > 0 and len(common_words) / len(para_words) > 0.8:
                    best_page = page_num
                    break
        
        # If found, assign to that page
        if best_page:
            current_page = best_page
            if best_page not in result:
                result[best_page] = []
            result[best_page].append(para_idx)
        else:
            # If not found, assign to current page (maintain sequence)
            if current_page not in result:
                result[current_page] = []
            result[current_page].append(para_idx)
    
    # Ensure all PDF pages are in result (even if empty)
    max_page = max(pdf_clean.keys())
    for page in range(1, max_page + 1):
        if page not in result:
            result[page] = []
    
    # Sort result by page number and paragraph indices within each page
    result = {k: sorted(v) for k, v in sorted(result.items())}
    
    return result